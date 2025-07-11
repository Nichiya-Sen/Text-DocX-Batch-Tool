# -*- coding: utf-8 -*-
import os
import tkinter as tk
from tkinter import filedialog, messagebox, scrolledtext
from tkinter import ttk
from tkinterdnd2 import TkinterDnD, DND_FILES
from docx import Document
import docx
from docx.text.paragraph import Paragraph as DocxParagraph
from docx.table import Table as DocxTable

import re
from PIL import Image, ImageTk
import customtkinter as ctk
import threading
import time
import copy
import traceback
import sys
sys.setrecursionlimit(3000)

from app_languages import LANGUAGES

# --- 外觀設定 ---
ctk.set_appearance_mode("System")
ctk.set_default_color_theme("blue")

# --- 字型定義 ---
BASE_FONT_SIZE = 14
APP_FONT = ("微軟正黑體", BASE_FONT_SIZE)
APP_FONT_BOLD = ("微軟正黑體", BASE_FONT_SIZE, "bold")
APP_FONT_LARGE_BOLD = ("微軟正黑體", BASE_FONT_SIZE + 2, "bold")
TOOLTIP_FONT = ("微軟正黑體", BASE_FONT_SIZE - 2)

# --- 圖示路徑定義 ---
# 根據是否為 PyInstaller 打包的程式，動態設定圖示基礎路徑
if getattr(sys, 'frozen', False) and hasattr(sys, '_MEIPASS'):
    # 如果是 PyInstaller 打包的程式，使用臨時資料夾路徑
    # PyInstaller 會將 --add-data 添加的資源解壓縮到 sys._MEIPASS 下
    # 這裡的 "icons" 是你在打包命令中 --add-data "icons;icons" 的第二個參數（目標路徑）
    ICON_BASE_PATH = os.path.join(sys._MEIPASS, "icons")
else:
    # 否則，使用原始的相對路徑 (用於開發環境)
    ICON_BASE_PATH = "icons" # 這裡也改為不帶斜線，os.path.join 會正確處理

# 確保所有圖示路徑使用這個 ICON_BASE_PATH
FOLDER_ICON_PATH = os.path.join(ICON_BASE_PATH, "folder_icon.png")
SEARCH_ICON_PATH = os.path.join(ICON_BASE_PATH, "search_icon.png")
UNDO_ICON_PATH = os.path.join(ICON_BASE_PATH, "undo_icon.png")
REDO_ICON_PATH = os.path.join(ICON_BASE_PATH, "redo_icon.png")
SAVE_ICON_PATH = os.path.join(ICON_BASE_PATH, "save_icon.png")
SAVE_ALL_ICON_PATH = os.path.join(ICON_BASE_PATH, "save_all_icon.png")
DELETE_ICON_PATH = os.path.join(ICON_BASE_PATH, "delete_icon.png")
BATCH_UNDO_ICON_PATH = os.path.join(ICON_BASE_PATH, "batch_undo_icon.png")
REFRESH_ICON_PATH = os.path.join(ICON_BASE_PATH, "refresh_icon.png")
HELP_ICON_PATH = os.path.join(ICON_BASE_PATH, "help_icon.png")
SETTINGS_ICON_PATH = os.path.join(ICON_BASE_PATH, "settings_icon.png")
MAGNIFY_ICON_PATH = os.path.join(ICON_BASE_PATH, "magnify_icon.png")
ADD_FILES_ICON_PATH = os.path.join(ICON_BASE_PATH, "add_files_icon.png") # 新增：匯入檔案圖示

# --- 圖示載入輔助函數 ---
def load_icon(icon_path, size):
    if os.path.exists(icon_path):
        try:
            icon_image = Image.open(icon_path)
            return ctk.CTkImage(light_image=icon_image, dark_image=icon_image, size=size)
        except Exception as e:
            print(f"警告：載入或處理圖示檔案失敗：{icon_path} - {e}")
            return None
    return None

# --- 其他全域常數 ---
BTN_ICON_SIZE = (int(BASE_FONT_SIZE * 1.8), int(BASE_FONT_SIZE * 1.8))
TREEVIEW_SCROLL_STEP = 10

# --- 自然排序輔助函數 ---
def natural_sort_key(s):
    return [int(text) if text.isdigit() else text.lower() for text in re.split('([0-9]+)', s)]

class Tooltip:
    def __init__(self, widget, text_key, delay=1000):
        self.widget = widget
        self.text_key = text_key
        self.delay = delay
        self.tooltip_window = None
        self.show_id = None

        self.widget.bind("<Enter>", self.schedule_show_tooltip, add="+")
        self.widget.bind("<Leave>", self.hide_tooltip, add="+")
        self.widget.bind("<ButtonPress>", self.hide_tooltip, add="+")

    def schedule_show_tooltip(self, event=None):
        self.hide_tooltip()

        app_instance = None
        current_widget = self.widget
        # 遍歷父元件以找到 WordEditorApp 實例 (通常是 root)
        while current_widget:
            if hasattr(current_widget, 'tr') and hasattr(current_widget, 'lang_code'):
                app_instance = current_widget
                break
            # 安全地獲取 master
            if hasattr(current_widget, 'master'):
                 current_master = getattr(current_widget, 'master')
                 if current_master is None: # 已達頂層
                     current_widget = None
                 else:
                     current_widget = current_master
            else: # 沒有 master 屬性
                 current_widget = None

        if app_instance and self.text_key:
            try:
                actual_text = app_instance.tr(self.text_key)
                if actual_text : # 確保 actual_text 不是 None 或空字串
                    self.show_id = self.widget.after(self.delay, lambda: self.show_tooltip_with_text(actual_text))
            except Exception as e:
                # 如果翻譯過程出錯，並且 text_key 本身是個可顯示的字串，就顯示 text_key
                # 這有助於在 app_languages.py 中缺少鍵時進行除錯
                if isinstance(self.text_key, str) and self.text_key :
                    self.show_id = self.widget.after(self.delay, lambda: self.show_tooltip_with_text(self.text_key))


    def show_tooltip_with_text(self, text_to_display):
        if self.tooltip_window or not text_to_display:
            return

        x = self.widget.winfo_rootx() + 20
        y = self.widget.winfo_rooty() + self.widget.winfo_height() + 5

        self.tooltip_window = tk.Toplevel(self.widget)
        self.tooltip_window.wm_overrideredirect(True)
        self.tooltip_window.wm_geometry(f"+{x}+{y}")

        label = tk.Label(self.tooltip_window, text=text_to_display, justify='left',
                         background="#FFFFE0", relief='solid', borderwidth=1,
                         font=TOOLTIP_FONT, wraplength=300)
        label.pack(ipadx=5, ipady=3)

    def hide_tooltip(self, event=None):
        if self.show_id:
            self.widget.after_cancel(self.show_id)
            self.show_id = None
        if self.tooltip_window:
            self.tooltip_window.destroy()
            self.tooltip_window = None

    def update_text_key(self, new_text_key):
        """當語言切換時，用來更新 Tooltip 的文字鍵名。"""
        self.text_key = new_text_key
        self.hide_tooltip() # 先隱藏當前的 Tooltip，下次滑鼠移入時會使用新的鍵名重新翻譯並顯示

class WordEditorApp:
    def __init__(self, root):
        self.root = root
        self.root.geometry("1200x750")
        self.root.minsize(width=1050, height=650)

        # 呼叫初始化函式來組織 __init__ 內容
        self._initialize_state()
        self._setup_styles_and_icons()
        self._setup_main_frames()
        self._setup_top_icon_buttons() # 頂部按鈕放在這裡
        self._setup_file_management_ui()
        self._setup_search_replace_ui()
        self._setup_center_editor_ui()
        self._setup_right_info_ui()
        self._bind_global_events()

        self.root.title(self.tr("app_title"))

        # 設定視窗圖示
        try:
            # 確保使用 PyInstaller 打包後的正確路徑
            if getattr(sys, 'frozen', False) and hasattr(sys, '_MEIPASS'):
                # 如果是打包後的程式，圖示會被放在 _MEIPASS/icons 資料夾
                icon_file_path = os.path.join(sys._MEIPASS, "icons", "app_icon.ico")
            else:
                # 如果是開發環境，圖示在當前目錄的 icons 資料夾
                icon_file_path = os.path.join("icons", "app_icon.ico")

            if os.path.exists(icon_file_path):
                self.root.iconbitmap(icon_file_path)
            else:
                print(f"警告：視窗圖示檔案未找到：{icon_file_path}")
        except Exception as e:
            print(f"設定視窗圖示時發生錯誤：{e}")

        self.update_file_tree_display()
        self.update_word_count()
        # 確保在按鈕被建立後再設定其初始狀態
        if hasattr(self, 'btn_undo_batch'):
             self.btn_undo_batch.configure(state="disabled")
        self.root.protocol("WM_DELETE_WINDOW", self.on_closing)
        self.on_tab_change() # Call this to set initial tab state and button states

    def _initialize_state(self):
        """初始化應用程式狀態相關的屬性。"""
        self.lang_code = "zh-tw" # 預設語言
        self.root.tr = self.tr # 讓根視窗也能存取翻譯函式，方便 Tooltip 使用
        self.root.lang_code = self.lang_code # 讓根視窗也能存取目前語言代碼

        self.files = []
        self.current_file = None
        self.current_file_type = None
        self.current_doc = None # 用於儲存目前開啟的 .docx Document 物件
        self.current_text_content = "" # 用於儲存目前檔案載入時的原始純文字內容
        self.search_results = {}
        self.modified_files = set() # 儲存已修改但未儲存的檔案路徑
        self.selected_files = {} # 儲存檔案列表中各檔案的勾選狀態 {file_path: True/False}
        self.modified_docs = {} # 儲存已修改的 .docx Document 物件 (用於格式化取代後、儲存前)
        self.modified_texts = {} # 儲存已修改的純文字內容 (用於txt，或docx純文字編輯後)
        self.file_timestamps = {} # 儲存檔案的最後修改時間戳記
        self.undo_stack = []
        self.redo_stack = []
        self.undo_limit = 100
        self._is_undoing_redoing = False # 標記是否正在執行復原/重做，避免觸發不必要的事件
        self.last_batch_originals = {} # 儲存上次批次取代前的原始檔案內容 (用於復原批次)
        self.text_editor_font_size = BASE_FONT_SIZE # 從 BASE_FONT_SIZE 初始化，這樣可以同步初始值
        self.sort_column = "filename" # 預設排序欄位
        self.sort_reverse = False # 預設排序方向 (False 為升冪)
        self._is_text_editor_enabled = False # 標記純文字編輯器是否可編輯

    def _setup_styles_and_icons(self):
        """設定 ttk 元件的樣式並載入所有圖示。"""
        # --- 檢查並建立圖示資料夾 ---
        if not os.path.exists(ICON_BASE_PATH):
            try:
                os.makedirs(ICON_BASE_PATH)
            except OSError as e:
                print(f"警告：無法建立圖示資料夾 {ICON_BASE_PATH} - {e}")

        # --- 載入所有圖示 ---
        self.folder_icon_img = load_icon(FOLDER_ICON_PATH, BTN_ICON_SIZE)
        self.search_icon_img = load_icon(SEARCH_ICON_PATH, BTN_ICON_SIZE)
        self.undo_icon_img = load_icon(UNDO_ICON_PATH, BTN_ICON_SIZE)
        self.redo_icon_img = load_icon(REDO_ICON_PATH, BTN_ICON_SIZE)
        self.save_icon_img = load_icon(SAVE_ICON_PATH, BTN_ICON_SIZE)
        self.save_all_icon_img = load_icon(SAVE_ALL_ICON_PATH, BTN_ICON_SIZE)
        self.delete_icon_img = load_icon(DELETE_ICON_PATH, BTN_ICON_SIZE)
        self.batch_undo_icon_img = load_icon(BATCH_UNDO_ICON_PATH, BTN_ICON_SIZE)
        self.refresh_icon_img = load_icon(REFRESH_ICON_PATH, BTN_ICON_SIZE)
        self.help_icon_img = load_icon(HELP_ICON_PATH, BTN_ICON_SIZE)
        self.settings_icon_img = load_icon(SETTINGS_ICON_PATH, BTN_ICON_SIZE)
        self.magnify_icon_img = load_icon(MAGNIFY_ICON_PATH, BTN_ICON_SIZE)
        self.add_files_icon_img = load_icon(ADD_FILES_ICON_PATH, BTN_ICON_SIZE) # 新增：匯入檔案圖示

        # --- 設定 ttk 元件的樣式 ---
        style = ttk.Style()
        style.configure("Treeview", font=APP_FONT, rowheight=int(BASE_FONT_SIZE * 2.0))
        style.configure("Treeview.Heading", font=APP_FONT_BOLD)
        style.configure("Vertical.TScrollbar", troughcolor=style.lookup("TFrame", "background"),
                        bordercolor=style.lookup("TFrame", "background"), arrowcolor=style.lookup("TFrame", "foreground"),
                        background=style.lookup("TFrame", "background"))
        style.configure("Horizontal.TScrollbar", troughcolor=style.lookup("TFrame", "background"),
                        bordercolor=style.lookup("TFrame", "background"), arrowcolor=style.lookup("TFrame", "foreground"),
                        background=style.lookup("TFrame", "background"))

        # 確保 Treeview 選擇的項目即使失去焦點也能保持藍色背景
        # CustomTkinter 的主題色可能會覆蓋，所以這裡嘗試更明確的設定
        style.map('Treeview',
            background=[('selected', '#3e84d4')], # 預設選擇時的背景色（藍色）
            foreground=[('selected', 'white')], # 選擇時的文字顏色
        )
        # 為了解決選取藍色框消失的問題，我們將在 on_file_tree_click 中手動應用tag
        style.configure("Custom.Treeview", font=APP_FONT, rowheight=int(BASE_FONT_SIZE * 2.0))
        style.map('Custom.Treeview',
            background=[('selected', '#3e84d4')],  # 預設選擇時的背景色（藍色）
            foreground=[('selected', 'white')],    # 選擇時的文字顏色
        )
        # 這裡不再需要 style.tag_configure，因為 tag_configure 是 Treeview 實例的方法


    def _setup_main_frames(self):
        """建立主視窗佈局框架。"""
        self.main_frame = ctk.CTkFrame(self.root, corner_radius=0, fg_color="transparent")
        self.main_frame.pack(padx=10, pady=10, fill="both", expand=True)

        self.main_frame.grid_rowconfigure(0, weight=1)
        self.main_frame.grid_columnconfigure(0, weight=0)
        self.main_frame.grid_columnconfigure(1, weight=1)
        self.main_frame.grid_columnconfigure(2, weight=0)

        self.left_frame = ctk.CTkFrame(self.main_frame, width=420, corner_radius=8)
        self.left_frame.grid(row=0, column=0, sticky="ns", padx=(0,5), pady=5)
        self.left_frame.grid_propagate(False)
        self.left_frame.grid_rowconfigure(0, weight=2)
        self.left_frame.grid_rowconfigure(1, weight=3)
        self.left_frame.grid_columnconfigure(0, weight=1)

        self.center_frame = ctk.CTkFrame(self.main_frame, corner_radius=8)
        self.center_frame.grid(row=0, column=1, sticky="nsew", padx=5, pady=5)
        # Center frame now has 2 rows for content + font slider row
        self.center_frame.grid_rowconfigure(0, weight=0) # Header
        self.center_frame.grid_rowconfigure(1, weight=1) # Tabview
        self.center_frame.grid_rowconfigure(2, weight=0) # Font slider
        self.center_frame.grid_columnconfigure(0, weight=1)

        self.right_frame = ctk.CTkFrame(self.main_frame, width=380, corner_radius=8)
        self.right_frame.grid(row=0, column=2, sticky="ns", padx=(5,0), pady=5)
        self.right_frame.grid_propagate(False)
        self.right_frame.grid_rowconfigure(1, weight=1)
        self.right_frame.grid_rowconfigure(3, weight=1)
        self.right_frame.grid_columnconfigure(0, weight=1)

    def _setup_top_icon_buttons(self):
        """設定頂部圖示按鈕 (設定與幫助)。"""
        self.top_icon_frame = ctk.CTkFrame(self.root, fg_color="transparent")
        self.top_icon_frame.pack(side="top", anchor="ne", fill="x", padx=8, pady=2)
        self.top_icon_frame.grid_columnconfigure(0, weight=1)
        self.top_icon_frame.grid_columnconfigure(1, weight=0)
        self.top_icon_frame.grid_columnconfigure(2, weight=0)

        self.app_settings_button = ctk.CTkButton(
            self.top_icon_frame,
            image=self.settings_icon_img,
            text="",
            width=BTN_ICON_SIZE[0],
            height=BTN_ICON_SIZE[1],
            command=self.open_settings_dialog
        )
        self.app_settings_button.grid(row=0, column=1, sticky="e", padx=(2, 2))
        self.tooltip_app_settings_button = Tooltip(self.app_settings_button, "settings_tooltip")

        self.app_help_button = ctk.CTkButton(
            self.top_icon_frame,
            image=self.help_icon_img,
            text="",
            width=BTN_ICON_SIZE[0],
            height=BTN_ICON_SIZE[1],
            command=self.open_help_dialog
        )
        self.app_help_button.grid(row=0, column=2, sticky="e", padx=(2, 0))
        self.tooltip_app_help_button = Tooltip(self.app_help_button, "help_tooltip")

    def _setup_file_management_ui(self):
        """設定左側的檔案管理界面。"""
        self.file_management_frame = ctk.CTkFrame(self.left_frame, fg_color="transparent")
        self.file_management_frame.grid(row=0, column=0, sticky="nsew", padx=5, pady=(5,0))
        self.file_management_frame.grid_columnconfigure(0, weight=1)
        self.file_management_frame.grid_rowconfigure(3, weight=1)

        self.frame_files = ctk.CTkFrame(self.file_management_frame, fg_color="transparent")
        self.label_folder = ctk.CTkLabel(self.frame_files, text=self.tr("folder_section_label"), font=APP_FONT)
        self.label_folder.pack(side="left", padx=(0,2), anchor="w")
        self.entry_folder = ctk.CTkEntry(self.frame_files, font=APP_FONT)
        self.entry_folder.pack(side="left", padx=2, fill="x", expand=True, anchor="center")

        if self.folder_icon_img:
            self.btn_browse = ctk.CTkButton(self.frame_files, image=self.folder_icon_img, text="", command=self.load_folder, width=BTN_ICON_SIZE[0], height=BTN_ICON_SIZE[1])
        else:
            self.btn_browse = ctk.CTkButton(self.frame_files, text=self.tr("browse_folder_button_text_no_icon"), command=self.load_folder, font=APP_FONT, width=int(BASE_FONT_SIZE*3.5))
        self.btn_browse.pack(side="left", padx=(2,0), anchor="e")
        self.tooltip_btn_browse = Tooltip(self.btn_browse, "browse_folder_tooltip")

        # 新增匯入檔案按鈕
        if self.add_files_icon_img:
            self.btn_add_files = ctk.CTkButton(self.frame_files, image=self.add_files_icon_img, text="", command=self.add_files, width=BTN_ICON_SIZE[0], height=BTN_ICON_SIZE[1])
        else:
            self.btn_add_files = ctk.CTkButton(self.frame_files, text=self.tr("add_files_button_text_no_icon"), command=self.add_files, font=APP_FONT, width=int(BASE_FONT_SIZE*4))
        self.btn_add_files.pack(side="left", padx=(2,0), anchor="e") # 放在瀏覽按鈕旁邊
        self.tooltip_btn_add_files = Tooltip(self.btn_add_files, "add_files_tooltip")

        self.frame_files.pack(fill="x", pady=(0,5)) # 這行確保 frame_files 被pack

        file_list_header_frame = ctk.CTkFrame(self.file_management_frame, fg_color="transparent")
        file_list_header_frame.pack(fill="x", pady=(2,2))
        self.label_files = ctk.CTkLabel(file_list_header_frame, text=self.tr("loaded_files_label"), font=APP_FONT_BOLD)
        self.label_files.pack(side="left", anchor="w")

        sort_button_frame = ctk.CTkFrame(file_list_header_frame, fg_color="transparent")
        sort_button_frame.pack(side="right", anchor="e")
        self.btn_sort_name = ctk.CTkButton(sort_button_frame, text=self.tr("sort_name_button"), command=lambda: self.sort_files("filename"), font=APP_FONT, width=50, height=int(BASE_FONT_SIZE*1.8))
        self.btn_sort_name.pack(side="left", padx=1)
        self.tooltip_btn_sort_name = Tooltip(self.btn_sort_name, "sort_name_tooltip")
        self.btn_sort_date = ctk.CTkButton(sort_button_frame, text=self.tr("sort_date_button"), command=lambda: self.sort_files("date"), font=APP_FONT, width=50, height=int(BASE_FONT_SIZE*1.8))
        self.btn_sort_date.pack(side="left", padx=1)
        self.tooltip_btn_sort_date = Tooltip(self.btn_sort_date, "sort_date_tooltip")

        checkbox_control_frame = ctk.CTkFrame(self.file_management_frame, fg_color="transparent")
        checkbox_control_frame.pack(fill="x", pady=(0,2))
        self.select_all_var = tk.BooleanVar(value=False)
        self.select_all_check = ctk.CTkCheckBox(checkbox_control_frame, text=self.tr("select_all_checkbox"), variable=self.select_all_var, command=self.toggle_select_all, font=APP_FONT)
        self.select_all_check.pack(side="left", padx=5)
        self.tooltip_select_all_check = Tooltip(self.select_all_check, "select_all_tooltip")

        if self.delete_icon_img:
             self.btn_delete_selected = ctk.CTkButton(checkbox_control_frame, image=self.delete_icon_img, text="", command=self.delete_selected_files, width=BTN_ICON_SIZE[0]+2, height=BTN_ICON_SIZE[1]+2)
        else:
             self.btn_delete_selected = ctk.CTkButton(checkbox_control_frame, text=self.tr("remove_selected_button_text_no_icon"), command=self.delete_selected_files, font=APP_FONT, width=70)
        self.btn_delete_selected.pack(side="right", padx=5)
        self.tooltip_btn_delete_selected = Tooltip(self.btn_delete_selected, "remove_selected_tooltip")

        file_tree_frame = ctk.CTkFrame(self.file_management_frame, fg_color="transparent")
        file_tree_frame.pack(fill="both", expand=True, pady=0)

        # 使用自定義樣式來解決藍色框問題
        self.file_tree = ttk.Treeview(file_tree_frame, columns=("checkbox", "filename"), show="headings", height=7, style="Custom.Treeview")
        self.file_tree.heading("checkbox", text=self.tr("file_tree_col_checkbox"), anchor="center")
        self.file_tree.heading("filename", text=self.tr("file_tree_col_filename"), anchor="w")
        self.file_tree.column("checkbox", width=35, minwidth=35, stretch=tk.NO, anchor="center")
        self.file_tree.column("filename", width=1000, minwidth=350, stretch=tk.NO, anchor="w")
        self.file_tree.grid(row=0, column=0, sticky="nsew")
        v_scroll_ft = ttk.Scrollbar(file_tree_frame, orient="vertical", command=self.file_tree.yview)
        v_scroll_ft.grid(row=0, column=1, sticky="ns")
        self.file_tree.configure(yscrollcommand=v_scroll_ft.set)
        h_scroll_ft = ttk.Scrollbar(file_tree_frame, orient="horizontal", command=self.file_tree.xview)
        h_scroll_ft.grid(row=1, column=0, sticky="ew")
        self.file_tree.configure(xscrollcommand=h_scroll_ft.set)
        file_tree_frame.columnconfigure(0, weight=1)
        file_tree_frame.rowconfigure(0, weight=1)
        self.file_tree.bind('<ButtonRelease-1>', self.on_file_tree_click)
        self.file_tree.bind('<Shift-MouseWheel>', self._on_shift_mouse_wheel_file_tree)
        self.file_tree.tag_configure("red", foreground="red")
        self.file_tree.tag_configure("black", foreground="black")
        self.file_tree.tag_configure("green", foreground="green") # 新增標籤用於搜尋結果
        self.file_tree.tag_configure("blue", foreground="blue") # 新增標籤用於搜尋結果
        # 將 tag_configure 移到 self.file_tree 建立之後
        self.file_tree.tag_configure("selected_item_tag", background="#3e84d4", foreground="white") # 新增用於選中項目標籤

    def _setup_search_replace_ui(self):
        """設定左側的搜尋/取代界面。"""
        self.search_replace_frame = ctk.CTkFrame(self.left_frame, fg_color="transparent")
        self.search_replace_frame.grid(row=1, column=0, sticky="nsew", padx=5, pady=0)
        self.search_replace_frame.grid_columnconfigure(0, weight=1)
        self.search_replace_frame.grid_rowconfigure(3, weight=1)

        self.label_search_replace_title = ctk.CTkLabel(self.search_replace_frame, text=self.tr("search_replace_title_label"), font=APP_FONT_BOLD)
        self.label_search_replace_title.pack(anchor="w", pady=(5,0), padx=5)
        self.auto_save_var = tk.BooleanVar(value=False)
        self.auto_save_check = ctk.CTkCheckBox(self.search_replace_frame, text=self.tr("auto_save_checkbox"), variable=self.auto_save_var, font=APP_FONT)
        self.auto_save_check.pack(anchor="w", pady=(2,2), padx=5)
        self.tooltip_auto_save_check = Tooltip(self.auto_save_check, "auto_save_tooltip")

        input_fields_main_frame = ctk.CTkFrame(self.search_replace_frame, fg_color="transparent")
        input_fields_main_frame.pack(fill="x", padx=5)
        search_field_frame = ctk.CTkFrame(input_fields_main_frame, fg_color="transparent")
        search_field_frame.pack(fill="x", pady=1)
        self.label_search_input = ctk.CTkLabel(search_field_frame, text=self.tr("search_input_label"), font=APP_FONT)
        self.label_search_input.pack(side="left", padx=(0,2))
        self.entry_search = ctk.CTkEntry(search_field_frame, font=APP_FONT)
        self.entry_search.pack(side="left", fill="x", expand=True)
        self.tooltip_entry_search = Tooltip(self.entry_search, "search_input_tooltip")

        replace_field_frame = ctk.CTkFrame(input_fields_main_frame, fg_color="transparent")
        replace_field_frame.pack(fill="x", pady=1)
        self.label_replace_input = ctk.CTkLabel(replace_field_frame, text=self.tr("replace_input_label"), font=APP_FONT)
        self.label_replace_input.pack(side="left", padx=(0,2))
        self.entry_replace = ctk.CTkEntry(replace_field_frame, font=APP_FONT)
        self.entry_replace.pack(side="left", fill="x", expand=True)
        self.tooltip_entry_replace = Tooltip(self.entry_replace, "replace_input_tooltip")

        self.replace_rules_label = ctk.CTkLabel(self.search_replace_frame, text=self.tr("replace_rules_list_label"), font=APP_FONT)
        self.replace_rules_label.pack(anchor="w", pady=(2,0), padx=5)
        replace_tree_container = ctk.CTkFrame(self.search_replace_frame, fg_color="transparent")
        replace_tree_container.pack(fill="both", expand=True, pady=(2,5), padx=5)

        self.replace_tree = ttk.Treeview(replace_tree_container, columns=("search", "replace"), show="headings", height=5)
        self.replace_tree.heading("search", text=self.tr("replace_tree_col_search"))
        self.replace_tree.heading("replace", text=self.tr("replace_tree_col_replace"))
        self.replace_tree.column("search", width=200, minwidth=100, stretch=tk.NO, anchor="w")
        self.replace_tree.column("replace", width=200, minwidth=100, stretch=tk.NO, anchor="w")
        self.replace_tree.grid(row=0, column=0, sticky="nsew")
        v_scroll_rt = ttk.Scrollbar(replace_tree_container, orient="vertical", command=self.replace_tree.yview)
        v_scroll_rt.grid(row=0, column=1, sticky="ns")
        self.replace_tree.configure(yscrollcommand=v_scroll_rt.set)
        h_scroll_rt = ttk.Scrollbar(replace_tree_container, orient="horizontal", command=self.replace_tree.xview)
        h_scroll_rt.grid(row=1, column=0, sticky="ew")
        self.replace_tree.configure(xscrollcommand=h_scroll_rt.set)
        replace_tree_container.grid_rowconfigure(0, weight=1)
        replace_tree_container.grid_columnconfigure(0, weight=1)
        self.replace_tree.bind('<Shift-MouseWheel>', self._on_shift_mouse_wheel_replace_tree)

        sr_btn_frame = ctk.CTkFrame(self.search_replace_frame, fg_color="transparent")
        sr_btn_frame.pack(fill="x", pady=(2,5), padx=5)

        self.btn_add_rule = ctk.CTkButton(sr_btn_frame, text=self.tr("add_rule_button"), command=self.add_replace_rule, font=APP_FONT)
        self.btn_add_rule.pack(side="left", padx=2, fill="x", expand=True)
        self.tooltip_btn_add_rule = Tooltip(self.btn_add_rule, "add_rule_tooltip")
        self.btn_remove_rule = ctk.CTkButton(sr_btn_frame, text=self.tr("remove_rule_button"), command=self.remove_replace_rule, font=APP_FONT)
        self.btn_remove_rule.pack(side="left", padx=2, fill="x", expand=True)
        self.tooltip_btn_remove_rule = Tooltip(self.btn_remove_rule, "remove_rule_tooltip")

        search_btn_actual_width = BTN_ICON_SIZE[0]+2 if self.search_icon_img else 50
        if self.search_icon_img:
            self.btn_search = ctk.CTkButton(sr_btn_frame, image=self.search_icon_img, text="", command=self.search_text, width=search_btn_actual_width, height=BTN_ICON_SIZE[1]+2)
        else:
            self.btn_search = ctk.CTkButton(sr_btn_frame, text=self.tr("search_button_text_no_icon"), command=self.search_text, font=APP_FONT, width=search_btn_actual_width)
        self.btn_search.pack(side="left", padx=2)
        self.tooltip_btn_search = Tooltip(self.btn_search, "search_button_tooltip")

        sr_btn_frame_2 = ctk.CTkFrame(self.search_replace_frame, fg_color="transparent")
        sr_btn_frame_2.pack(fill="x", pady=(0,5), padx=5)

        self.btn_replace_text_edit = ctk.CTkButton(sr_btn_frame_2, text=self.tr("plain_text_batch_button"), command=lambda: self.execute_replace_action("plain_text_batch"), font=APP_FONT)
        self.btn_replace_text_edit.pack(side="left", padx=2, fill="x", expand=True)
        self.tooltip_btn_replace_text_edit = Tooltip(self.btn_replace_text_edit, "plain_text_batch_tooltip")

        self.btn_replace_formatted = ctk.CTkButton(sr_btn_frame_2, text=self.tr("formatted_batch_button"), command=lambda: self.execute_replace_action("formatted_batch"), font=APP_FONT)
        self.btn_replace_formatted.pack(side="left", padx=2, fill="x", expand=True)
        self.tooltip_btn_replace_formatted = Tooltip(self.btn_replace_formatted, "formatted_batch_tooltip")

        undo_batch_actual_width = BTN_ICON_SIZE[0]+2 if self.batch_undo_icon_img else 70
        if self.batch_undo_icon_img:
            self.btn_undo_batch = ctk.CTkButton(sr_btn_frame_2, image=self.batch_undo_icon_img, text="", command=self.undo_last_batch_replace, width=undo_batch_actual_width, height=BTN_ICON_SIZE[1]+2, state="disabled")
        else:
            self.btn_undo_batch = ctk.CTkButton(sr_btn_frame_2, text=self.tr("undo_batch_button_text_no_icon"), command=self.undo_last_batch_replace, font=APP_FONT, width=undo_batch_actual_width, state="disabled")
        self.btn_undo_batch.pack(side="left", padx=2)
        self.tooltip_btn_undo_batch = Tooltip(self.btn_undo_batch, "undo_batch_tooltip")

    def _setup_center_editor_ui(self):
        """設定中間的編輯器和字體大小控制。"""
        header_frame_center = ctk.CTkFrame(self.center_frame, fg_color="transparent")
        header_frame_center.grid(row=0, column=0, sticky="ew", pady=(5,2), padx=5)
        header_frame_center.grid_columnconfigure(0, weight=1)
        header_frame_center.grid_columnconfigure(1, weight=0)
        header_frame_center.grid_columnconfigure(2, weight=0)
        header_frame_center.grid_columnconfigure(3, weight=0)

        self.label_current_file = ctk.CTkLabel(header_frame_center, text=f"{self.tr('current_file_label_prefix')} {self.tr('current_file_label_not_loaded')}",  font=APP_FONT_LARGE_BOLD, anchor="w")
        self.label_current_file.grid(row=0, column=0, sticky="ew", padx=(0,5))

        self.label_word_count = ctk.CTkLabel(header_frame_center, text=f"{self.tr('word_count_label_prefix')} {self.tr('word_count_label_none')}", font=APP_FONT, anchor="e")
        self.label_word_count.grid(row=0, column=1, sticky="e", padx=(5,5))

        refresh_btn_actual_width = BTN_ICON_SIZE[0]+2 if self.refresh_icon_img else int(BASE_FONT_SIZE*3.5)
        if self.refresh_icon_img:
            self.btn_refresh_current = ctk.CTkButton(header_frame_center, image=self.refresh_icon_img, text="", command=self.reload_current_file_from_disk, width=refresh_btn_actual_width, height=BTN_ICON_SIZE[1]+2)
        else:
            self.btn_refresh_current = ctk.CTkButton(header_frame_center, text=self.tr("refresh_current_button_text_no_icon"), command=self.reload_current_file_from_disk, font=APP_FONT, width=refresh_btn_actual_width)
        self.btn_refresh_current.grid(row=0, column=2, sticky="e", padx=(0,5))
        self.tooltip_btn_refresh_current = Tooltip(self.btn_refresh_current, "refresh_current_tooltip")

        button_row_frame_center = ctk.CTkFrame(header_frame_center, fg_color="transparent")
        button_row_frame_center.grid(row=0, column=3, sticky="e")
        btn_op_width = BTN_ICON_SIZE[0]+2 if self.undo_icon_img else int(BASE_FONT_SIZE*3.5)

        if self.undo_icon_img:
             self.btn_undo = ctk.CTkButton(button_row_frame_center, image=self.undo_icon_img, text="", command=self.undo, width=btn_op_width, height=BTN_ICON_SIZE[1]+2)
        else:
             self.btn_undo = ctk.CTkButton(button_row_frame_center, text=self.tr("undo_button_text_no_icon"), command=self.undo, font=APP_FONT, width=btn_op_width)
        self.btn_undo.pack(side="left", padx=(0,2))
        self.tooltip_btn_undo = Tooltip(self.btn_undo, "undo_tooltip")

        if self.redo_icon_img:
             self.btn_redo = ctk.CTkButton(button_row_frame_center, image=self.redo_icon_img, text="", command=self.redo, width=btn_op_width, height=BTN_ICON_SIZE[1]+2)
        else:
             self.btn_redo = ctk.CTkButton(button_row_frame_center, text=self.tr("redo_button_text_no_icon"), command=self.redo, font=APP_FONT, width=btn_op_width)
        self.btn_redo.pack(side="left", padx=2)
        self.tooltip_btn_redo = Tooltip(self.btn_redo, "redo_tooltip")

        if self.save_icon_img:
             self.btn_save_current = ctk.CTkButton(button_row_frame_center, image=self.save_icon_img, text="", command=self.save_file, width=btn_op_width, height=BTN_ICON_SIZE[1]+2)
        else:
             self.btn_save_current = ctk.CTkButton(button_row_frame_center, text=self.tr("save_current_button_text_no_icon"), command=self.save_file, font=APP_FONT, width=btn_op_width)
        self.btn_save_current.pack(side="left", padx=2)
        self.tooltip_btn_save_current = Tooltip(self.btn_save_current, "save_current_tooltip")

        if self.save_all_icon_img:
             self.btn_save_all = ctk.CTkButton(button_row_frame_center, image=self.save_all_icon_img, text="", command=self.save_all_selected_files, width=btn_op_width, height=BTN_ICON_SIZE[1]+2)
        else:
             self.btn_save_all = ctk.CTkButton(button_row_frame_center, text=self.tr("save_all_button_text_no_icon"), command=self.save_all_selected_files, font=APP_FONT, width=btn_op_width)
        self.btn_save_all.pack(side="left", padx=(2,0))
        self.tooltip_btn_save_all = Tooltip(self.btn_save_all, "save_all_tooltip")

        # --- Tabview (initialized in _create_tab_view method) ---
        self.tab_view = None
        self.text_content = None
        self.formatted_replace_content_frame = None
        self.label_formatted_replace_info = None
        self.formatted_preview_text = None
        self._create_tab_view() # Call a separate method to create tabview and its contents

        # --- Font Size Control Frame ---
        self.font_control_frame = ctk.CTkFrame(self.center_frame, fg_color="transparent")
        self.font_control_frame.grid(row=2, column=0, sticky="ew", padx=10, pady=(0, 10))
        self.font_control_frame.grid_columnconfigure(1, weight=1)

        if self.magnify_icon_img:
            self.magnify_label = ctk.CTkLabel(self.font_control_frame, text="", image=self.magnify_icon_img)
            self.magnify_label.grid(row=0, column=0, padx=(0,5))
        else:
            self.magnify_label = ctk.CTkLabel(self.font_control_frame, text="Aa", font=APP_FONT_BOLD)
            self.magnify_label.grid(row=0, column=0, padx=(0,5))

        self.font_size_slider = ctk.CTkSlider(self.font_control_frame, from_=8, to=72, number_of_steps=64, command=self.update_font_size_from_slider)
        self.font_size_slider.set(self.text_editor_font_size) # Set initial position based on current font size
        self.font_size_slider.grid(row=0, column=1, sticky="ew", padx=(0,5))

        # 將 label_font_size 改為 CTkEntry
        self.font_size_entry_var = tk.StringVar(value=str(self.text_editor_font_size))
        self.font_size_entry = ctk.CTkEntry(
            self.font_control_frame,
            textvariable=self.font_size_entry_var,
            width=50,
            font=APP_FONT,
            justify='center' # 讓文字置中
        )
        self.font_size_entry.grid(row=0, column=2, padx=(5,0))
        self.font_size_entry.bind('<Return>', self._on_font_size_entry_change)
        self.font_size_entry.bind('<FocusOut>', self._on_font_size_entry_change)


    def _setup_right_info_ui(self):
        """設定右側的搜尋結果和訊息列表。"""
        self.label_results = ctk.CTkLabel(self.right_frame, text=self.tr("search_results_label"), font=APP_FONT_BOLD)
        self.label_results.grid(row=0, column=0, sticky="ew", pady=(5,0), padx=5)
        tree_frame_right = ctk.CTkFrame(self.right_frame, fg_color="transparent")
        tree_frame_right.grid(row=1, column=0, sticky="nsew", pady=(0,2), padx=5)
        tree_frame_right.grid_columnconfigure(0, weight=1)
        tree_frame_right.grid_rowconfigure(0, weight=1)

        self.tree = ttk.Treeview(tree_frame_right, columns=("filename", "count"), show="headings", height=10)
        self.tree.heading("filename", text=self.tr("search_tree_col_filename"))
        self.tree.heading("count", text=self.tr("search_tree_col_count"))
        self.tree.column("filename", width=1000, minwidth=250, stretch=tk.NO, anchor="w")
        self.tree.column("count", width=60, minwidth=40, stretch=tk.NO, anchor="center")
        self.tree.grid(row=0, column=0, sticky="nsew")
        v_scroll_results = ttk.Scrollbar(tree_frame_right, orient="vertical", command=self.tree.yview)
        v_scroll_results.grid(row=0, column=1, sticky="ns")
        self.tree.configure(yscrollcommand=v_scroll_results.set)
        h_scroll_results = ttk.Scrollbar(tree_frame_right, orient="horizontal", command=self.tree.xview)
        h_scroll_results.grid(row=1, column=0, sticky="ew")
        self.tree.configure(xscrollcommand=h_scroll_results.set)
        self.tree.bind('<ButtonRelease-1>', self.on_tree_click)
        self.tree.bind('<Shift-MouseWheel>', self._on_shift_mouse_wheel_tree)
        self.tree.tag_configure("red", foreground="red")
        self.tree.tag_configure("black", foreground="black")

        self.label_messages = ctk.CTkLabel(self.right_frame, text=self.tr("messages_label"), font=APP_FONT_BOLD)
        self.label_messages.grid(row=2, column=0, sticky="ew", pady=(2,0), padx=5)
        message_frame_right = ctk.CTkFrame(self.right_frame, fg_color="transparent")
        message_frame_right.grid(row=3, column=0, sticky="nsew", pady=(0,5), padx=5)
        message_frame_right.grid_columnconfigure(0, weight=1)
        message_frame_right.grid_rowconfigure(0, weight=1)

        self.message_tree = ttk.Treeview(message_frame_right, columns=("message",), show="headings", height=5)
        self.message_tree.heading("message", text=self.tr("message_tree_col_message"))
        self.message_tree.column("message", width=1000, minwidth=250, stretch=tk.NO, anchor="w")
        self.message_tree.grid(row=0, column=0, sticky="nsew")
        v_scroll_messages = ttk.Scrollbar(message_frame_right, orient="vertical", command=self.message_tree.yview)
        v_scroll_messages.grid(row=0, column=1, sticky="ns")
        self.message_tree.configure(yscrollcommand=v_scroll_messages.set)
        h_scroll_messages = ttk.Scrollbar(message_frame_right, orient="horizontal", command=self.message_tree.xview)
        h_scroll_messages.grid(row=1, column=0, sticky="ew")
        self.message_tree.configure(xscrollcommand=h_scroll_messages.set)
        self.message_tree.bind('<Shift-MouseWheel>', self._on_shift_mouse_wheel_message_tree)

    def _bind_global_events(self):
        """綁定全域快捷鍵和拖放事件。"""
        self.root.bind('<Control-z>', self.undo)
        self.root.bind('<Control-Key-Z>', self.undo)
        self.root.bind('<Control-y>', self.redo)
        self.root.bind('<Control-Key-Y>', self.redo)
        self.root.bind('<Control-s>', lambda event: self.save_file())
        self.root.bind('<Control-S>', lambda event: self.save_file())
        self.root.bind('<Control-a>', lambda event: self.select_all(event))
        self.root.bind('<Control-f>', lambda event: self.focus_search())

        if hasattr(self, 'entry_search') and self.entry_search:
            self.entry_search.bind('<Return>', lambda event: self.search_text())
            self.entry_search.bind('<KP_Enter>', lambda event: self.search_text())

        self.root.drop_target_register(DND_FILES)
        self.root.dnd_bind('<<Drop>>', self.drop_files)

    def _create_tab_view(self):
        # Store current content to restore it after recreating tabs
        current_text_content_val = ""
        current_formatted_preview_text_val = ""
        current_tab_selected_name = ""

        # Check if text_content exists and is a valid widget before trying to get its content
        if hasattr(self, 'text_content') and self.text_content and self.text_content.winfo_exists():
            current_text_content_val = self.text_content.get("1.0", tk.END + "-1c")
        if hasattr(self, 'formatted_preview_text') and self.formatted_preview_text and self.formatted_preview_text.winfo_exists():
            current_formatted_preview_text_val = self.formatted_preview_text.get("1.0", tk.END + "-1c")
        if hasattr(self, 'tab_view') and self.tab_view and self.tab_view.winfo_exists():
            current_tab_selected_name = self.tab_view.get()
            self.tab_view.destroy() # Destroy existing tabview to recreate

        self.tab_view = ctk.CTkTabview(self.center_frame, command=self.on_tab_change)
        if hasattr(self.tab_view, "_segmented_button") and self.tab_view._segmented_button is not None:
            self.tab_view._segmented_button.configure(font=APP_FONT)
        else:
            print(f"{self.tr('warning_title')}: {self.tr('ctk_tabview_font_error')}")
        self.tab_view.grid(row=1, column=0, sticky="nsew", padx=5, pady=(0,5))

        self.tab_text_edit_page = self.tab_view.add(self.tr("tab_text_edit"))
        self.tab_formatted_replace_page = self.tab_view.add(self.tr("tab_formatted_replace"))

        # Recreate textboxes and bind events
        self.text_content = ctk.CTkTextbox(self.tab_text_edit_page, wrap="word", font=("微軟正黑體", self.text_editor_font_size), state="disabled")
        self.text_content.pack(fill="both", expand=True, padx=0, pady=0)
        self.text_content.bind("<KeyRelease>", self.on_text_content_key_release)
        self.text_content.bind("<<Paste>>", self.on_text_content_key_release)
        self.text_content.bind("<Delete>", self.on_text_content_key_release)
        self.text_content.bind("<BackSpace>", self.on_text_content_key_release)
        # Fix: Ensure event is passed to undo/redo/save/select_all methods
        self.text_content.bind('<Control-y>', self.redo)
        self.text_content.bind('<Control-Y>', self.redo)
        self.text_content.bind('<Control-z>', self.undo)
        self.text_content.bind('<Control-Z>', self.undo)
        self.text_content.bind('<Control-s>', lambda event: self.save_file(event=event)) # Pass event
        self.text_content.bind('<Control-S>', lambda event: self.save_file(event=event)) # Pass event
        self.text_content.bind('<Control-a>', self.select_all_text)
        self.text_content.bind('<Control-MouseWheel>', self._on_ctrl_mouse_wheel)
        self.text_content.bind('<Control-Button-4>', self._on_ctrl_mouse_wheel)
        self.text_content.bind('<Control-Button-5>', self._on_ctrl_mouse_wheel)
        self.text_content.bind('<Control-plus>', self._on_ctrl_plus_minus)
        self.text_content.bind('<Control-minus>', self._on_ctrl_plus_minus)
        self.text_content.bind('<Control-KP_Add>', self._on_ctrl_plus_minus)
        self.text_content.bind('<Control-KP_Subtract>', self._on_ctrl_plus_minus)
        self.text_content.tag_config("highlight", foreground="red")

        self.formatted_replace_content_frame = ctk.CTkFrame(self.tab_formatted_replace_page, fg_color="transparent")
        self.formatted_replace_content_frame.pack(fill="both", expand=True, padx=0, pady=0)

        self.label_formatted_replace_info = ctk.CTkLabel(self.formatted_replace_content_frame,
                                                         text=self.tr("formatted_replace_info_label"),
                                                         font=APP_FONT, wraplength=450, justify="left")
        self.label_formatted_replace_info.pack(padx=20, pady=(10, 5), anchor="w")

        self.formatted_preview_text = ctk.CTkTextbox(self.formatted_replace_content_frame, wrap="word", font=("微軟正黑體", self.text_editor_font_size), state="disabled")
        self.formatted_preview_text.pack(fill="both", expand=True, padx=20, pady=(0, 20))
        self.formatted_preview_text.tag_config("highlight", foreground="red")
        self.formatted_preview_text.bind('<Control-MouseWheel>', self._on_ctrl_mouse_wheel)
        self.formatted_preview_text.bind('<Control-Button-4>', self._on_ctrl_mouse_wheel)
        self.formatted_preview_text.bind('<Control-Button-5>', self._on_ctrl_mouse_wheel)
        self.formatted_preview_text.bind('<Control-plus>', self._on_ctrl_plus_minus)
        self.formatted_preview_text.bind('<Control-minus>', self._on_ctrl_plus_minus)
        self.formatted_preview_text.bind('<Control-KP_Add>', self._on_ctrl_plus_minus)
        self.formatted_preview_text.bind('<Control-KP_Subtract>', self._on_ctrl_plus_minus)

        # Restore content and set selected tab
        if current_text_content_val:
            self.text_content.configure(state="normal") # Temporarily enable to insert
            self.text_content.insert("1.0", current_text_content_val)
            self.text_content.configure(state="disabled") # Restore state based on logic in on_tab_change
        if current_formatted_preview_text_val:
            self.formatted_preview_text.configure(state="normal") # Temporarily enable to insert
            self.formatted_preview_text.insert("1.0", current_formatted_preview_text_val)
            self.formatted_preview_text.configure(state="disabled") # Restore state based on logic in on_tab_change

        # Determine the target tab name based on current language
        target_tab_name = ""
        # Prioritize matching by current translated name first
        if current_tab_selected_name == self.tr("tab_text_edit"):
            target_tab_name = self.tr("tab_text_edit")
        elif current_tab_selected_name == self.tr("tab_formatted_replace"):
            target_tab_name = self.tr("tab_formatted_replace")
        # Fallback to English internal keys if the exact translated name isn't found (e.g., first run after language change)
        elif current_tab_selected_name == "Plain Text Edit":
            target_tab_name = self.tr("tab_text_edit")
        elif current_tab_selected_name == "Formatted Replace":
            target_tab_name = self.tr("tab_formatted_replace")
        else: # Default to first tab if no specific tab was selected or old name is unrecognized
            if self.tab_view._segmented_button.cget("values"):
                target_tab_name = self.tab_view._segmented_button.cget("values")[0]


        if target_tab_name and target_tab_name in self.tab_view._tab_dict:
            self.tab_view.set(target_tab_name)
        else:
            if self.tab_view._segmented_button.cget("values"):
                self.tab_view.set(self.tab_view._segmented_button.cget("values")[0])

    def update_font_size_from_slider(self, value):
        new_size = int(value)
        if new_size != self.text_editor_font_size:
            self.text_editor_font_size = new_size
            font_name = APP_FONT[0]
            if hasattr(self, 'text_content') and self.text_content.winfo_exists():
                self.text_content.configure(font=(font_name, self.text_editor_font_size))
            if hasattr(self, 'formatted_preview_text') and self.formatted_preview_text.winfo_exists():
                self.formatted_preview_text.configure(font=(font_name, self.text_editor_font_size))
            # 同步更新輸入框的顯示值
            if hasattr(self, 'font_size_entry_var'):
                self.font_size_entry_var.set(str(self.text_editor_font_size))

    def _on_font_size_entry_change(self, event=None):
        try:
            new_size_str = self.font_size_entry_var.get().strip()
            new_size = int(new_size_str)
            # 限制字體大小在合理範圍內
            new_size = max(8, min(72, new_size))
            if new_size != self.text_editor_font_size:
                self.text_editor_font_size = new_size
                font_name = APP_FONT[0]
                if hasattr(self, 'text_content') and self.text_content.winfo_exists():
                    self.text_content.configure(font=(font_name, self.text_editor_font_size))
                if hasattr(self, 'formatted_preview_text') and self.formatted_preview_text.winfo_exists():
                    self.formatted_preview_text.configure(font=(font_name, self.text_editor_font_size))
                # 同步更新 slider 的位置
                if hasattr(self, 'font_size_slider'):
                    self.font_size_slider.set(new_size)
            # 確保輸入框的值與實際字體大小一致 (可能因為限制範圍而改變)
            self.font_size_entry_var.set(str(self.text_editor_font_size))
        except ValueError:
            # 如果輸入無效，恢復到當前的字體大小
            if hasattr(self, 'font_size_entry_var'):
                self.font_size_entry_var.set(str(self.text_editor_font_size))


    def toggle_select_all(self):
        select_all_status = self.select_all_var.get()

        for file_path in self.files:
            self.selected_files[file_path] = select_all_status

        self.update_file_tree_display()

    def _on_shift_mouse_wheel_file_tree(self, event):
        if hasattr(self, 'file_tree') and self.file_tree.winfo_exists():
            if event.delta > 0: self.file_tree.xview_scroll(-TREEVIEW_SCROLL_STEP, "units")
            else: self.file_tree.xview_scroll(TREEVIEW_SCROLL_STEP, "units")
            return "break"

    def _on_shift_mouse_wheel_replace_tree(self, event):
        if hasattr(self, 'replace_tree') and self.replace_tree.winfo_exists():
            if event.delta > 0: self.replace_tree.xview_scroll(-TREEVIEW_SCROLL_STEP, "units")
            else: self.replace_tree.xview_scroll(TREEVIEW_SCROLL_STEP, "units")
            return "break"

    def _on_shift_mouse_wheel_tree(self, event):
        if hasattr(self, 'tree') and self.tree.winfo_exists():
            if event.delta > 0: self.tree.xview_scroll(-TREEVIEW_SCROLL_STEP, "units")
            else: self.tree.xview_scroll(TREEVIEW_SCROLL_STEP, "units")
            return "break"

    def _on_shift_mouse_wheel_message_tree(self, event):
        if hasattr(self, 'message_tree') and self.message_tree.winfo_exists():
            if event.delta > 0: self.message_tree.xview_scroll(-TREEVIEW_SCROLL_STEP, "units")
            else: self.message_tree.xview_scroll(TREEVIEW_SCROLL_STEP, "units")
            return "break"

    def _on_ctrl_mouse_wheel(self, event):
        if hasattr(self, 'font_size_slider') and self.font_size_slider.winfo_exists():
            current_value = self.font_size_slider.get()
            if event.delta > 0 or event.num == 4:
                new_value = min(self.font_size_slider.cget("to"), current_value + 1)
            else:
                new_value = max(self.font_size_slider.cget("from_"), current_value - 1)
            self.font_size_slider.set(new_value)
            self.update_font_size_from_slider(new_value)
        return "break"

    def _on_ctrl_plus_minus(self, event):
        if hasattr(self, 'font_size_slider') and self.font_size_slider.winfo_exists():
            current_value = self.font_size_slider.get()
            if event.keysym == "plus" or event.keysym == "KP_Add":
                new_value = min(self.font_size_slider.cget("to"), current_value + 1)
            elif event.keysym == "minus" or event.keysym == "KP_Subtract":
                new_value = max(self.font_size_slider.cget("from_"), current_value - 1)
            else:
                return
            self.font_size_slider.set(new_value)
            self.update_font_size_from_slider(new_value)
        return "break"

    def select_all_text(self, event=None):
        focused_widget = self.root.focus_get()
        if isinstance(focused_widget, ctk.CTkTextbox):
            focused_widget.tag_add(tk.SEL, "1.0", tk.END)
            focused_widget.mark_set(tk.INSERT, "1.0")
            focused_widget.see(tk.INSERT)
            return 'break'
        elif isinstance(focused_widget, ctk.CTkEntry):
            focused_widget.select_range(0, tk.END)
            return 'break'
        return None

    def focus_search(self):
        if hasattr(self, 'entry_search') and self.entry_search.winfo_exists():
            self.entry_search.focus_set()
            self.entry_search.select_range(0, tk.END)

    def update_select_all_state(self):
        if not (hasattr(self, 'select_all_var') and hasattr(self, 'files') and self.files): return
        all_selected = True
        any_selected = False
        if not self.files: all_selected = False
        else:
            for f_path in self.files:
                if self.selected_files.get(f_path, False): any_selected = True
                else: all_selected = False
        self.select_all_var.set(all_selected)

    def _highlight_active_editor(self, search_term=None):
        if search_term is None:
            search_term_to_use = self.entry_search.get() if hasattr(self, 'entry_search') else ""
        else:
            search_term_to_use = search_term

        if hasattr(self, 'text_content') and self.text_content.winfo_exists():
            self.text_content.tag_remove("highlight", "1.0", tk.END)
        if hasattr(self, 'formatted_preview_text') and self.formatted_preview_text.winfo_exists():
            self.formatted_preview_text.tag_remove("highlight", "1.0", tk.END)

        if not search_term_to_use or not self.current_file:
            return

        editor_to_highlight = None
        current_tab_name = self.tab_view.get() if hasattr(self, 'tab_view') else ""

        if current_tab_name == self.tr("tab_text_edit") and \
           hasattr(self, 'text_content') and self.text_content.winfo_exists():
            editor_to_highlight = self.text_content
        elif current_tab_name == self.tr("tab_formatted_replace") and \
             hasattr(self, 'formatted_preview_text') and self.formatted_preview_text.winfo_exists():
             editor_to_highlight = self.formatted_preview_text

        if not editor_to_highlight:
            return

        start_pos = "1.0"
        while True:
            match_var = tk.IntVar()
            start_pos = editor_to_highlight.search(search_term_to_use, start_pos, tk.END, nocase=True, count=match_var)
            if not start_pos or match_var.get() == 0:
                break
            end_pos = f"{start_pos}+{match_var.get()}c"
            editor_to_highlight.tag_add("highlight", start_pos, end_pos)
            start_pos = end_pos

    def _clear_text_editor(self):
        if hasattr(self, 'text_content') and self.text_content.winfo_exists():
            self.text_content.configure(state="normal")
            self.text_content.delete("1.0", tk.END)
            self.text_content.tag_remove("highlight", "1.0", tk.END)

    def _clear_formatted_preview(self):
        if hasattr(self, 'formatted_preview_text') and self.formatted_preview_text.winfo_exists():
            self.formatted_preview_text.configure(state="normal")
            self.formatted_preview_text.delete("1.0", tk.END)
            self.formatted_preview_text.tag_remove("highlight", "1.0", tk.END)
            self.formatted_preview_text.configure(state="disabled")

    def on_text_content_key_release(self, event=None):
        if self._is_undoing_redoing or not self.current_file or not self._is_text_editor_enabled:
             return

        current_editor_content = self.text_content.get("1.0", tk.END).strip()

        previous_content_for_compare = self.current_text_content.strip()
        if self.current_file in self.modified_texts:
             previous_content_for_compare = self.modified_texts[self.current_file].strip()

        if current_editor_content != previous_content_for_compare:
            self.modified_files.add(self.current_file)
            self.modified_texts[self.current_file] = current_editor_content
            self._save_text_content_state(current_editor_content)
        self.update_word_count()
        self.update_file_tree_display()
        self._highlight_active_editor(self.entry_search.get())
        self._update_undo_redo_buttons() # 在內容改變時更新按鈕狀態

    def _save_text_content_state(self, content):
        if not self.undo_stack or self.undo_stack[-1] != content:
            self.undo_stack.append(content)
            if len(self.undo_stack) > self.undo_limit:
                self.undo_stack.pop(0)
            self.redo_stack = []
            self._update_undo_redo_buttons() # 狀態改變時更新按鈕

    def undo(self, event=None):
        if len(self.undo_stack) > 1 and self.current_file and self._is_text_editor_enabled:
            self._is_undoing_redoing = True
            last_state = self.undo_stack.pop()
            self.redo_stack.append(last_state)
            content_to_restore = self.undo_stack[-1]

            current_pos = self.text_content.index(tk.INSERT)
            self.text_content.configure(state="normal")
            self.text_content.delete("1.0", tk.END)
            self.text_content.insert("1.0", content_to_restore)
            try: self.text_content.mark_set(tk.INSERT, current_pos)
            except tk.TclError: self.text_content.mark_set(tk.INSERT, tk.END)

            self.modified_texts[self.current_file] = content_to_restore
            self.modified_files.add(self.current_file)
            self.update_word_count()
            self.update_file_tree_display()
            self._is_undoing_redoing = False
            self._highlight_active_editor(self.entry_search.get())
            self._update_undo_redo_buttons() # 執行後更新按鈕
        return "break"


    def redo(self, event=None):
        if self.redo_stack and self.current_file and self._is_text_editor_enabled:
            self._is_undoing_redoing = True
            content_to_restore = self.redo_stack.pop()
            self.undo_stack.append(content_to_restore)

            current_pos = self.text_content.index(tk.INSERT)
            self.text_content.configure(state="normal")
            self.text_content.delete("1.0", tk.END)
            self.text_content.insert("1.0", content_to_restore)
            try: self.text_content.mark_set(tk.INSERT, current_pos)
            except tk.TclError: self.text_content.mark_set(tk.INSERT, tk.END)

            self.modified_texts[self.current_file] = content_to_restore
            self.modified_files.add(self.current_file)
            self.update_word_count()
            self.update_file_tree_display()
            self._is_undoing_redoing = False
            self._highlight_active_editor(self.entry_search.get())
            self._update_undo_redo_buttons() # 執行後更新按鈕
        return "break"

    def _update_undo_redo_buttons(self):
        """更新復原和重做按鈕的啟用/禁用狀態。"""
        if hasattr(self, 'btn_undo'):
            # 只有當純文字編輯器啟用且有可復原的內容時，才啟用復原按鈕
            self.btn_undo.configure(state="normal" if self._is_text_editor_enabled and len(self.undo_stack) > 1 else "disabled")
        if hasattr(self, 'btn_redo'):
            # 只有當純文字編輯器啟用且有可重做的內容時，才啟用重做按鈕
            self.btn_redo.configure(state="normal" if self._is_text_editor_enabled and self.redo_stack else "disabled")


    def on_tab_change(self, event=None):
        if not hasattr(self, 'tab_view') or not hasattr(self, 'text_content') or \
           not hasattr(self, 'formatted_preview_text'):
            return

        active_tab_display_name = self.tab_view.get()
        is_file_loaded = bool(self.current_file)

        if active_tab_display_name == self.tr("tab_text_edit"):
            self.text_content.configure(state="normal" if is_file_loaded else "disabled")
            self._is_text_editor_enabled = is_file_loaded
            # 這裡不直接設定 undo/redo 按鈕狀態，而是呼叫統一的更新函數
            self._update_undo_redo_buttons()
        else:
            self.text_content.configure(state="disabled")
            self._is_text_editor_enabled = False
            # 這裡不直接設定 undo/redo 按鈕狀態，而是呼叫統一的更新函數
            self._update_undo_redo_buttons()


        if active_tab_display_name == self.tr("tab_formatted_replace") and is_file_loaded:
            preview_content = self.modified_texts.get(self.current_file, self.current_text_content)
            self.formatted_preview_text.configure(state="normal")
            self.formatted_preview_text.delete("1.0", tk.END)
            self.formatted_preview_text.insert("1.0", preview_content)
            self.formatted_preview_text.configure(state="disabled")
        else:
            self.formatted_preview_text.configure(state="disabled")

        if is_file_loaded:
            self.btn_replace_text_edit.configure(state="normal" if active_tab_display_name == self.tr("tab_text_edit") else "disabled")
            self.btn_replace_formatted.configure(state="normal" if active_tab_display_name == self.tr("tab_formatted_replace") else "disabled")
        else:
            self.btn_replace_text_edit.configure(state="disabled")
            self.btn_replace_formatted.configure(state="disabled")

        self.update_word_count()
        self._highlight_active_editor(self.entry_search.get())

    def tr(self, key, **kwargs):
        current_lang_texts = LANGUAGES.get(self.lang_code)

        if current_lang_texts is None:
            default_lang_code = "zh-tw"
            if default_lang_code not in LANGUAGES and LANGUAGES:
                default_lang_code = list(LANGUAGES.keys())[0]
            current_lang_texts = LANGUAGES.get(default_lang_code, {})

        translation = current_lang_texts.get(key, key)

        if kwargs:
            try:
                return translation.format(**kwargs)
            except Exception:
                return translation
        return translation

    def open_help_dialog(self):
        help_win = ctk.CTkToplevel(self.root)
        help_win.title(self.tr("help_dialog_title"))
        help_win.geometry("600x450")
        help_win.resizable(True, True)

        help_text_area = ctk.CTkTextbox(help_win, wrap="word", font=APP_FONT)
        help_text_area.pack(padx=10, pady=10, fill="both", expand=True)

        # 修正這裡：明確組合 help_content 和 help_author_info
        help_content_text = self.tr("help_content") + "\n" + \
                            self.tr("help_author_info") + "\n" + \
                            self.tr("help_completion_date")
        help_text_area.insert("1.0", help_content_text)
        help_text_area.configure(state="disabled")

        button_frame = ctk.CTkFrame(help_win, fg_color="transparent")
        button_frame.pack(pady=(0, 10))

        ok_button = ctk.CTkButton(button_frame, text=self.tr("ok_button"), command=help_win.destroy, font=APP_FONT)
        ok_button.pack()
        help_win.after(100, help_win.lift)
        help_win.grab_set()

    def open_settings_dialog(self):
        settings_win = ctk.CTkToplevel(self.root)
        settings_win.title(self.tr("settings_dialog_title"))
        settings_win.geometry("320x200")
        settings_win.resizable(False, False)

        main_dialog_frame = ctk.CTkFrame(settings_win, fg_color="transparent")
        main_dialog_frame.pack(padx=20, pady=20, fill="both", expand=True)

        label = ctk.CTkLabel(main_dialog_frame, text=self.tr("language_select_label"), font=APP_FONT)
        label.pack(pady=(0, 10), anchor="w")

        lang_var = tk.StringVar(value=self.lang_code)

        radio_button_container = ctk.CTkFrame(main_dialog_frame, fg_color="transparent")
        radio_button_container.pack(fill="x", expand=True, pady=(0,10))

        for code in LANGUAGES.keys():
            rb = ctk.CTkRadioButton(
                radio_button_container,
                text=self.tr(f"lang_{code}"),
                variable=lang_var,
                value=code,
                font=APP_FONT
            )
            rb.pack(anchor="w", padx=10, pady=3)

        def apply_lang_and_close():
            new_lang_code = lang_var.get()
            if self.lang_code != new_lang_code:
                self.lang_code = new_lang_code
                if hasattr(self.root, 'lang_code'):
                    self.root.lang_code = self.lang_code
                self.refresh_ui_texts()
            settings_win.destroy()

        apply_button = ctk.CTkButton(main_dialog_frame, text=self.tr("ok_button"), command=apply_lang_and_close, font=APP_FONT)
        apply_button.pack(pady=(10, 0), anchor="e")
        settings_win.after(100, settings_win.lift)
        settings_win.grab_set()

    def refresh_ui_texts(self):
        self.root.title(self.tr("app_title"))

        if hasattr(self, 'tooltip_app_help_button'): self.tooltip_app_help_button.update_text_key("help_tooltip")
        if hasattr(self, 'tooltip_app_settings_button'): self.tooltip_app_settings_button.update_text_key("settings_tooltip")

        if hasattr(self, 'label_folder'): self.label_folder.configure(text=self.tr("folder_section_label"))
        if hasattr(self, 'btn_browse'):
            if not self.folder_icon_img: self.btn_browse.configure(text=self.tr("browse_folder_button_text_no_icon"))
            if hasattr(self, 'tooltip_btn_browse'): self.tooltip_btn_browse.update_text_key("browse_folder_tooltip")
        # 更新新增檔案按鈕的文字和提示
        if hasattr(self, 'btn_add_files'):
            if not self.add_files_icon_img: self.btn_add_files.configure(text=self.tr("add_files_button_text_no_icon"))
            if hasattr(self, 'tooltip_btn_add_files'): self.tooltip_btn_add_files.update_text_key("add_files_tooltip")

        if hasattr(self, 'label_files'): self.label_files.configure(text=self.tr("loaded_files_label"))
        if hasattr(self, 'btn_sort_name'):
             self.btn_sort_name.configure(text=self.tr("sort_name_button"))
             if hasattr(self, 'tooltip_btn_sort_name'): self.tooltip_btn_sort_name.update_text_key("sort_name_tooltip")
        if hasattr(self, 'btn_sort_date'):
             self.btn_sort_date.configure(text=self.tr("sort_date_button"))
             if hasattr(self, 'tooltip_btn_sort_date'): self.tooltip_btn_sort_date.update_text_key("sort_date_tooltip")
        if hasattr(self, 'select_all_check'):
             self.select_all_check.configure(text=self.tr("select_all_checkbox"))
             if hasattr(self, 'tooltip_select_all_check'): self.tooltip_select_all_check.update_text_key("select_all_tooltip")
        if hasattr(self, 'btn_delete_selected'):
            if not self.delete_icon_img: self.btn_delete_selected.configure(text=self.tr("remove_selected_button_text_no_icon"))
            if hasattr(self, 'tooltip_btn_delete_selected'): self.tooltip_btn_delete_selected.update_text_key("remove_selected_tooltip")
        if hasattr(self, 'file_tree'):
            self.file_tree.heading("checkbox", text=self.tr("file_tree_col_checkbox"))
            self.file_tree.heading("filename", text=self.tr("file_tree_col_filename"))

        if hasattr(self, 'label_search_replace_title'): self.label_search_replace_title.configure(text=self.tr("search_replace_title_label"))
        if hasattr(self, 'auto_save_check'):
            self.auto_save_check.configure(text=self.tr("auto_save_checkbox"))
            if hasattr(self, 'tooltip_auto_save_check'): self.tooltip_auto_save_check.update_text_key("auto_save_tooltip")
        if hasattr(self, 'label_search_input'): self.label_search_input.configure(text=self.tr("search_input_label"))
        if hasattr(self, 'entry_search') and hasattr(self, 'tooltip_entry_search'): self.tooltip_entry_search.update_text_key("search_input_tooltip")
        if hasattr(self, 'label_replace_input'): self.label_replace_input.configure(text=self.tr("replace_input_label"))
        if hasattr(self, 'entry_replace') and hasattr(self, 'tooltip_entry_replace'): self.tooltip_entry_replace.update_text_key("replace_input_tooltip")
        if hasattr(self, 'replace_rules_label'): self.replace_rules_label.configure(text=self.tr("replace_rules_list_label"))
        if hasattr(self, 'replace_tree'):
            self.replace_tree.heading("search", text=self.tr("replace_tree_col_search"))
            self.replace_tree.heading("replace", text=self.tr("replace_tree_col_replace"))
        if hasattr(self, 'btn_add_rule'):
            self.btn_add_rule.configure(text=self.tr("add_rule_button"))
            if hasattr(self, 'tooltip_btn_add_rule'): self.tooltip_btn_add_rule.update_text_key("add_rule_tooltip")
        if hasattr(self, 'btn_remove_rule'):
            self.btn_remove_rule.configure(text=self.tr("remove_rule_button"))
            if hasattr(self, 'tooltip_btn_remove_rule'): self.tooltip_btn_remove_rule.update_text_key("remove_rule_tooltip")
        if hasattr(self, 'btn_search'):
            if not self.search_icon_img: self.btn_search.configure(text=self.tr("search_button_text_no_icon"))
            if hasattr(self, 'tooltip_btn_search'): self.tooltip_btn_search.update_text_key("search_button_tooltip")
        if hasattr(self, 'btn_replace_text_edit'):
            self.btn_replace_text_edit.configure(text=self.tr("plain_text_batch_button"))
            if hasattr(self, 'tooltip_btn_replace_text_edit'): self.tooltip_btn_replace_text_edit.update_text_key("plain_text_batch_tooltip")
        if hasattr(self, 'btn_replace_formatted'):
            self.btn_replace_formatted.configure(text=self.tr("formatted_batch_button"))
            if hasattr(self, 'tooltip_btn_replace_formatted'): self.tooltip_btn_replace_formatted.update_text_key("formatted_batch_tooltip")
        if hasattr(self, 'btn_undo_batch'):
            if not self.batch_undo_icon_img: self.btn_undo_batch.configure(text=self.tr("undo_batch_button_text_no_icon"))
            if hasattr(self, 'tooltip_btn_undo_batch'): self.tooltip_btn_undo_batch.update_text_key("undo_batch_tooltip")

        self.update_word_count()
        if hasattr(self, 'btn_refresh_current'):
            if not self.refresh_icon_img: self.btn_refresh_current.configure(text=self.tr("refresh_current_button_text_no_icon"))
            if hasattr(self, 'tooltip_btn_refresh_current'): self.tooltip_btn_refresh_current.update_text_key("refresh_current_tooltip")

        # 刷新按鈕狀態
        self._update_undo_redo_buttons()

        if hasattr(self, 'btn_save_current'):
            if not self.save_icon_img: self.btn_save_current.configure(text=self.tr("save_current_button_text_no_icon"))
            if hasattr(self, 'tooltip_btn_save_current'): self.tooltip_btn_save_current.update_text_key("save_current_tooltip")
        if hasattr(self, 'btn_save_all'):
            if not self.save_all_icon_img: self.btn_save_all.configure(text=self.tr("save_all_button_text_no_icon"))
            if hasattr(self, 'tooltip_btn_save_all'): self.tooltip_btn_save_all.update_text_key("save_all_tooltip")

        # Recreate tabview to ensure correct localized names are used as internal keys
        self._create_tab_view()

        # Re-apply font size to newly created textboxes
        self.update_font_size_from_slider(self.text_editor_font_size)

        if hasattr(self, 'label_formatted_replace_info'): self.label_formatted_replace_info.configure(text=self.tr("formatted_replace_info_label"))

        if hasattr(self, 'label_results'): self.label_results.configure(text=self.tr("search_results_label"))
        if hasattr(self, 'tree'):
            self.tree.heading("filename", text=self.tr("search_tree_col_filename"))
            self.tree.heading("count", text=self.tr("search_tree_col_count"))
        if hasattr(self, 'label_messages'): self.label_messages.configure(text=self.tr("messages_label"))
        if hasattr(self, 'message_tree'):
            self.message_tree.heading("message", text=self.tr("message_tree_col_message"))

        # 更新字體大小輸入框的顯示
        if hasattr(self, 'font_size_entry_var'):
            self.font_size_entry_var.set(str(self.text_editor_font_size))

        self.root.update_idletasks()

    def update_word_count(self):
        if not (hasattr(self, 'label_word_count') and self.label_word_count.winfo_exists() and
                hasattr(self, 'label_current_file') and self.label_current_file.winfo_exists()):
            return

        content_to_count = ""
        if self.current_file:
            active_tab_display_name = ""
            if hasattr(self, 'tab_view') and self.tab_view.get():
                active_tab_display_name = self.tab_view.get()

            if active_tab_display_name == self.tr("tab_text_edit") and self._is_text_editor_enabled and \
               hasattr(self, 'text_content') and self.text_content.winfo_exists():
                content_to_count = self.text_content.get("1.0", tk.END).strip()
            elif active_tab_display_name == self.tr("tab_formatted_replace") and \
                 hasattr(self, 'formatted_preview_text') and self.formatted_preview_text.winfo_exists():
                 content_to_count = self.formatted_preview_text.get("1.0", tk.END).strip()
            elif self.current_file:
                 content_to_count = self.modified_texts.get(self.current_file, self.current_text_content).strip()

        char_count = len(content_to_count) if content_to_count else 0

        current_file_display_text = self.tr("current_file_label_not_loaded")
        if self.current_file:
            base_filename = os.path.basename(self.current_file)
            current_file_display_text = f"*{base_filename}" if self.current_file in self.modified_files else base_filename

        self.label_current_file.configure(text=f"{self.tr('current_file_label_prefix')} {current_file_display_text}")

        word_count_display_text = self.tr("word_count_label_none")
        if self.current_file :
             word_count_display_text = f"{char_count}"
        self.label_word_count.configure(text=f"{self.tr('word_count_label_prefix')} {word_count_display_text}")

    def on_closing(self):
        if self.modified_files:
            unsaved_files_basenames = [os.path.basename(f) for f in self.modified_files]

            file_list_str = "\n".join(unsaved_files_basenames[:10])
            if len(unsaved_files_basenames) > 10:
                file_list_str += "\n..."

            # 從 app_languages.py 載入翻譯字串
            msg_template = self.tr("unsaved_changes_on_close_message")
            formatted_msg = msg_template.format(count=len(unsaved_files_basenames), file_list=file_list_str)

            if messagebox.askyesno(self.tr("unsaved_changes_on_close_title"), formatted_msg):
                self.root.destroy()
        else:
            self.root.destroy()

    def _add_message(self, message_key_or_text, **kwargs):
        if hasattr(self, 'message_tree') and self.message_tree.winfo_exists():
            timestamp = time.strftime("%H:%M:%S", time.localtime())

            display_message = ""
            # 優先嘗試作為翻譯鍵處理
            if isinstance(message_key_or_text, str) and message_key_or_text in LANGUAGES.get(self.lang_code, {}): # Check in current language dict
                translated_text = self.tr(message_key_or_text, **kwargs)
                display_message = translated_text
            # 如果不是翻譯鍵，但有 kwargs，嘗試作為格式化字串處理
            elif isinstance(message_key_or_text, str) and kwargs:
                try:
                    display_message = message_key_or_text.format(**kwargs)
                except (KeyError, ValueError): # 如果格式化失敗，就直接顯示原始字串
                    display_message = message_key_or_text
            # 如果只是純字串
            elif isinstance(message_key_or_text, str):
                 display_message = message_key_or_text
            # 其他情況，轉為字串
            else:
                display_message = str(message_key_or_text)

            self.message_tree.insert("", "end", values=(f"[{timestamp}] {display_message}",))
            self.message_tree.yview_moveto(1)
            self.root.after_idle(self.message_tree.update_idletasks)

    def load_folder(self, folder_path=None):
        if folder_path is None:
            folder_path = filedialog.askdirectory(title=self.tr("ask_directory_title"))
        if not folder_path:
            return

        self.entry_folder.delete(0, tk.END)
        self.entry_folder.insert(0, folder_path)
        current_open_file_path = self.current_file

        self.files = []
        self.selected_files = {}
        self.search_results = {}
        self.modified_files = set()
        self.modified_docs = {}
        self.modified_texts = {}
        self.file_timestamps = {}
        self.current_file = None
        self.current_file_type = None
        self.current_doc = None
        self.current_text_content = ""
        self._clear_text_editor()
        self._clear_formatted_preview()
        self.update_word_count()
        self.undo_stack = []
        self.redo_stack = []
        if hasattr(self, 'btn_undo_batch'): self.btn_undo_batch.configure(state="disabled")
        self.last_batch_originals.clear()

        loaded_count = 0
        try:
            for filename in os.listdir(folder_path):
                file_path = os.path.join(folder_path, filename)
                if os.path.isfile(file_path):
                    ext = os.path.splitext(filename)[1].lower().lstrip('.')
                    if ext in ['docx', 'txt'] and not filename.startswith('~$'):
                        self.files.append(file_path)
                        self.selected_files[file_path] = False
                        try:
                            self.file_timestamps[file_path] = os.path.getmtime(file_path)
                        except Exception as e_ts:
                             print(f"無法取得檔案時間戳記 {file_path}: {e_ts}")
                             self.file_timestamps[file_path] = 0
                        loaded_count += 1
        except Exception as e_list_dir:
            messagebox.showerror(self.tr("error_title"), self.tr("folder_read_error", folder_path=folder_path, error=str(e_list_dir)))
            self._add_message("folder_read_error_msg", folder_path=folder_path)
            self.files = []
            loaded_count = 0

        self.sort_files(self.sort_column, self.sort_reverse)
        self.update_file_tree_display()
        self.update_select_all_state()
        self.on_tab_change()
        self._update_undo_redo_buttons() # 載入文件夾時，更新撤銷/重做按鈕狀態

        if loaded_count > 0:
            self._add_message("files_loaded_from_folder_msg", count=loaded_count)
            if current_open_file_path and current_open_file_path in self.files:
                for item_id in self.file_tree.get_children():
                    if self.file_tree.set(item_id, "filename") == current_open_file_path:
                        self.file_tree.selection_set(item_id)
                        self.file_tree.focus(item_id)
                        self.load_file_content(current_open_file_path)
                        break
        else:
            if folder_path:
                self._add_message("no_supported_files_in_folder_msg", folder_name=os.path.basename(folder_path))

    def add_files(self):
        file_paths = filedialog.askopenfilenames(
            title=self.tr("ask_files_title"),
            # 修正 filetypes 順序和描述，以確保 TXT 檔案可見或作為預設選項
            filetypes=[(self.tr("text_documents"), "*.txt"),  # 將 TXT 放在最前面
                       (self.tr("word_documents"), "*.docx"),
                       (self.tr("all_files"), "*.*")]
        )
        if not file_paths:
            return

        added_count = 0
        for item_path in file_paths:
            file_path = item_path.strip('{}')
            if os.path.isfile(file_path):
                ext = os.path.splitext(file_path)[1].lower().lstrip('.')
                if ext in ['docx', 'txt'] and not os.path.basename(file_path).startswith('~$'):
                    if file_path not in self.files:
                        self.files.append(file_path)
                        self.selected_files[file_path] = False
                        try:
                            self.file_timestamps[file_path] = os.path.getmtime(file_path)
                        except Exception as e_ts:
                            print(f"警告：載入檔案時間戳記失敗 {file_path}: {e_ts}")
                            self.file_timestamps[file_path] = 0
                        added_count += 1
                    else:
                        self._add_message("file_already_in_list_msg", filename=os.path.basename(file_path))
                else:
                    self._add_message("unsupported_file_type_msg", filename=os.path.basename(file_path))
            else:
                self._add_message("invalid_file_path_msg", filepath=file_path)

        if added_count > 0:
            self.sort_files(self.sort_column, self.sort_reverse)
            self.update_file_tree_display()
            self.update_select_all_state()
            self._add_message("files_added_msg", count=added_count)
            self._update_undo_redo_buttons()
        elif file_paths:
            self._add_message("no_supported_files_selected_msg")


    def drop_files(self, event):
        file_paths_str = self.root.tk.splitlist(event.data)
        files_to_add_paths = []
        folders_to_process = [] # <--- 這個列表用於儲存要處理的資料夾
        for item_path_str in file_paths_str:
            item_path = item_path_str.strip('{}')
            if os.path.isdir(item_path):
                folders_to_process.append(item_path) # 修正：直接添加到 folders_to_process
            elif os.path.isfile(item_path):
                files_to_add_paths.append(item_path)
        added_count_total = 0
        for folder_path in folders_to_process: # 這裡現在可以正確地遍歷已添加的資料夾
            try:
                for filename in os.listdir(folder_path):
                    file_path = os.path.join(folder_path, filename)
                    if os.path.isfile(file_path):
                        ext = os.path.splitext(filename)[1].lower().lstrip('.')
                        if ext in ['docx', 'txt'] and not filename.startswith('~$'):
                            if file_path not in self.files:
                                self.files.append(file_path)
                                self.selected_files[file_path] = False
                                try: self.file_timestamps[file_path] = os.path.getmtime(file_path)
                                except: self.file_timestamps[file_path] = 0
                                added_count_total += 1
            except Exception as e_folder_drop:
                self._add_message("folder_read_error_msg", folder_path=os.path.basename(folder_path))
                print(f"拖曳資料夾 '{os.path.basename(folder_path)}' 時發生錯誤: {e_folder_drop}")
        for file_path in files_to_add_paths:
            ext = os.path.splitext(file_path)[1].lower().lstrip('.')
            if ext in ['docx', 'txt'] and not os.path.basename(file_path).startswith('~$'):
                if file_path not in self.files:
                    self.files.append(file_path)
                    self.selected_files[file_path] = False
                    try: self.file_timestamps[file_path] = os.path.getmtime(file_path)
                    except: self.file_timestamps[file_path] = 0
                    added_count_total += 1
                else:
                     self._add_message("file_already_in_list_msg", filename=os.path.basename(file_path))
        if added_count_total > 0:
            self.sort_files(self.sort_column, self.sort_reverse)
            self.update_file_tree_display()
            self.update_select_all_state()
            self._add_message("drag_drop_added_msg", count=added_count_total)
            self._update_undo_redo_buttons() # 拖放文件時，更新撤銷/重做按鈕狀態
        elif file_paths_str:
             self._add_message("drag_drop_no_valid_items_msg")

    def sort_files(self, column, reverse=None):
        if column == self.sort_column and reverse is None:
            self.sort_reverse = not self.sort_reverse
        elif reverse is not None:
            self.sort_reverse = reverse
        else:
            self.sort_reverse = False
        self.sort_column = column

        if column == "filename":
            self.files.sort(key=lambda f: natural_sort_key(os.path.basename(f)), reverse=self.sort_reverse)
        elif column == "date":
            try:
                self.files.sort(key=lambda f: os.path.getmtime(f) if os.path.exists(f) else 0, reverse=self.sort_reverse)
            except FileNotFoundError:
                 print("排序時發生 FileNotFoundError，部分檔案可能已被刪除。")
        self.update_file_tree_display()

    def update_file_tree_display(self):
        if not (hasattr(self, 'file_tree') and self.file_tree.winfo_exists()): return

        # 獲取當前選中的項目
        current_focused_item = self.file_tree.focus() # 直接獲取焦點項目的iid

        self.file_tree.delete(*self.file_tree.get_children())

        for file_path in self.files:
            filename = os.path.basename(file_path)
            display_filename = f"*{filename}" if file_path in self.modified_files else filename
            checkbox_state = "☑" if self.selected_files.get(file_path, False) else "☐"

            tags = []
            # 優先處理修改狀態的標籤
            if file_path in self.modified_files:
                tags.append("red")

            # 再處理搜尋結果的標籤
            if file_path in self.search_results and self.search_results[file_path] > 0 :
                if "red" not in tags:
                     tags.append("blue") # 藍色表示有搜尋結果但未修改
                else:
                    tags.append("green") # 綠色表示既修改又有搜尋結果

            # 如果這個文件是當前選中的文件，應用選中的標籤
            if file_path == self.current_file:
                tags.append("selected_item_tag") # 應用深藍色標籤

            try:
                # 插入項目並應用所有標籤
                self.file_tree.insert("", "end", values=(checkbox_state, display_filename), iid=file_path, tags=tuple(tags))
                # 將完整的路徑儲存為隱藏值
                self.file_tree.set(file_path, "filename", file_path)
            except tk.TclError as e:
                 if "duplicate item name" in str(e).lower():
                     print(f"警告：嘗試插入重複的 Treeview item ID: {file_path} - {e}")
                 else:
                     raise e

        # 重新選中之前有焦點的項目（如果它還存在）
        if current_focused_item and self.file_tree.exists(current_focused_item):
            self.file_tree.selection_set(current_focused_item)
            self.file_tree.focus(current_focused_item)
            # 確保選中的項目仍然應用 selected_item_tag
            current_tags = list(self.file_tree.item(current_focused_item, 'tags'))
            if "selected_item_tag" not in current_tags:
                current_tags.append("selected_item_tag")
                self.file_tree.item(current_focused_item, tags=tuple(current_tags))
            self.file_tree.see(current_focused_item) # 確保選中的項目可見


    def add_replace_rule(self):
        if not (hasattr(self, 'entry_search') and hasattr(self, 'entry_replace') and hasattr(self, 'replace_tree')):
            messagebox.showwarning(self.tr("warning_title"), self.tr("components_not_initialized_error"))
            return
        search_text = self.entry_search.get()
        replace_text = self.entry_replace.get()
        if not search_text:
            messagebox.showwarning(self.tr("warning_title"), self.tr("search_text_empty_warning"))
            return
        for item_id in self.replace_tree.get_children():
            values = self.replace_tree.item(item_id)['values']
            if values and len(values) > 0 and values[0] == search_text:
                messagebox.showwarning(self.tr("warning_title"), self.tr("search_text_already_in_rules_warning", search_text=search_text))
                return
        self.replace_tree.insert("", "end", values=(search_text, replace_text))
        if hasattr(self, 'replace_tree') and self.replace_tree.winfo_exists():
            self.root.after_idle(self.replace_tree.update_idletasks)
        self.entry_search.delete(0, tk.END)
        self.entry_replace.delete(0, tk.END)

    def remove_replace_rule(self):
        if not (hasattr(self, 'replace_tree') and self.replace_tree.winfo_exists()):
            messagebox.showwarning(self.tr("warning_title"), self.tr("components_not_initialized_error"))
            return
        selected_items = self.replace_tree.selection()
        if not selected_items:
            # 修改提示訊息，讓使用者更清楚
            messagebox.showwarning(self.tr("info_title"), self.tr("select_rule_to_delete_info"))
            return
        for item in selected_items:
            self.replace_tree.delete(item)
        if hasattr(self, 'replace_tree') and self.replace_tree.winfo_exists():
            self.root.after_idle(self.replace_tree.update_idletasks)

    def on_file_tree_click(self, event):
        if not (hasattr(self, 'file_tree') and self.file_tree.winfo_exists()): return

        # 清除所有項目上的 'selected_item_tag'
        for item_id in self.file_tree.get_children():
            current_tags = list(self.file_tree.item(item_id, 'tags'))
            if "selected_item_tag" in current_tags:
                current_tags.remove("selected_item_tag")
                self.file_tree.item(item_id, tags=tuple(current_tags))

        item_id = self.file_tree.identify_row(event.y)
        if not item_id or not self.file_tree.exists(item_id): return
        column_id = self.file_tree.identify_column(event.x)

        full_file_path = self.file_tree.set(item_id, "filename")
        try:
            displayed_values = self.file_tree.item(item_id, "values")
            displayed_filename_in_tree = displayed_values[1]
        except (IndexError, TypeError):
            return

        if column_id == "#1":
            current_state = self.selected_files.get(full_file_path, False)
            new_state = not current_state
            self.selected_files[full_file_path] = new_state
            checkbox_state_char = "☑" if new_state else "☐"
            self.file_tree.item(item_id, values=(checkbox_state_char, displayed_filename_in_tree))
            self.update_select_all_state()
        elif column_id == "#2":
            # 選中該行並應用 selected_item_tag
            self.file_tree.selection_set(item_id)
            self.file_tree.focus(item_id)
            current_tags = list(self.file_tree.item(item_id, 'tags'))
            if "selected_item_tag" not in current_tags:
                current_tags.append("selected_item_tag")
                self.file_tree.item(item_id, tags=tuple(current_tags))

            if self.current_file == full_file_path: return

            proceed_with_switch = True
            if self.current_file and self.current_file in self.modified_files:
                if self.auto_save_var.get():
                    self._add_message("auto_saving_modified_file_msg", filename=os.path.basename(self.current_file))
                    save_success = self.save_file(file_path_to_save=self.current_file)
                    if not save_success: proceed_with_switch = False
                else:
                    confirm_save = messagebox.askyesnocancel(
                        self.tr("save_changes_on_switch_title"),
                        self.tr("save_changes_on_switch_message", filename=os.path.basename(self.current_file))
                    )
                    if confirm_save is True:
                        save_success = self.save_file()
                        if not save_success: proceed_with_switch = False
                    elif confirm_save is False:
                        self.modified_files.discard(self.current_file)
                        self.modified_docs.pop(self.current_file, None)
                        self.modified_texts.pop(self.current_file, None)
                        self.update_file_tree_display()
                    else:
                        proceed_with_switch = False

                if not proceed_with_switch:
                    # 如果取消切換，重新選中之前的檔案
                    for prev_item_id in self.file_tree.get_children():
                        if self.file_tree.exists(prev_item_id) and self.file_tree.set(prev_item_id, "filename") == self.current_file:
                            if self.file_tree.exists(prev_item_id):
                                self.file_tree.selection_set(prev_item_id)
                                self.file_tree.focus(prev_item_id)
                                # 確保之前選中的檔案也重新應用 selected_item_tag
                                prev_tags = list(self.file_tree.item(prev_item_id, 'tags'))
                                if "selected_item_tag" not in prev_tags:
                                    prev_tags.append("selected_item_tag")
                                    self.file_tree.item(prev_item_id, tags=tuple(prev_tags))
                            break
                    return
            if proceed_with_switch:
                self.load_file_content(full_file_path)

    def load_file_content(self, file_path):
        if not os.path.exists(file_path):
            self._handle_file_not_exist_on_load(file_path)
            return

        if self.current_file == file_path and file_path in self.file_timestamps:
            try:
                current_mod_time = os.path.getmtime(file_path)
                # 這裡檢查是否在外部被修改，且當前檔案在應用程式中未被標記為修改狀態
                # 或者雖然被標記為修改，但編輯器未啟用（表示修改不是通過手動編輯引起的，可能是批次處理）
                # 避免頻繁彈窗
                if current_mod_time > self.file_timestamps.get(file_path, 0) and \
                   (file_path not in self.modified_files or not self._is_text_editor_enabled) :

                    msg_title = self.tr("external_file_change_prompt_title")
                    unsaved_info = self.tr("external_file_change_prompt_unsaved_app") if file_path in self.modified_files else ""
                    msg_body = self.tr("external_file_change_prompt_message",
                                       filename=os.path.basename(file_path),
                                       unsaved_changes_info=unsaved_info)

                    if messagebox.askyesno(msg_title, msg_body):
                        # 如果用戶選擇重新載入，則丟棄應用程式中的修改狀態
                        self.modified_files.discard(file_path)
                        self.modified_docs.pop(file_path, None)
                        self.modified_texts.pop(file_path, None)
                    else:
                        # 如果用戶選擇不重新載入，則更新時間戳以避免再次提示，並保持當前應用程式中的狀態
                        self.file_timestamps[file_path] = current_mod_time
                        self._add_message("user_chose_not_to_reload_msg", filename=os.path.basename(file_path))
                        return
            except FileNotFoundError:
                 self._handle_file_not_exist_on_load(file_path)
                 return
            except OSError as e:
                 print(f"檢查時間戳時發生 OS 錯誤: {file_path} - {e}")

        # 載入新文件時，清除舊的選取狀態
        if self.current_file:
            for item_id in self.file_tree.get_children():
                if self.file_tree.exists(item_id) and self.file_tree.set(item_id, "filename") == self.current_file:
                    current_tags = list(self.file_tree.item(item_id, 'tags'))
                    if "selected_item_tag" in current_tags:
                        current_tags.remove("selected_item_tag")
                        self.file_tree.item(item_id, tags=tuple(current_tags))
                    break
        self.current_file = file_path
        self.current_file_type = os.path.splitext(file_path)[1].lower().lstrip('.')
        self.current_doc = None
        self.current_text_content = ""

        self._clear_text_editor()
        self._clear_formatted_preview()
        self.text_content.configure(state="disabled")
        self.formatted_preview_text.configure(state="disabled")
        self._is_text_editor_enabled = False

        self.label_current_file.configure(text=f"{self.tr('current_file_label_prefix')} {self.tr('current_file_label_loading')}")
        self.label_word_count.configure(text=f"{self.tr('word_count_label_prefix')} {self.tr('word_count_label_none')}")

        threading.Thread(target=self._perform_file_load_in_background, args=(file_path,), daemon=True).start()

    def _handle_file_not_exist_on_load(self, file_path):
        messagebox.showerror(self.tr("error_title"), self.tr("file_not_exist_on_load_error", filename=os.path.basename(file_path)))
        if file_path in self.files: self.files.remove(file_path)
        self.selected_files.pop(file_path, None)
        self.modified_files.discard(file_path)
        self.modified_docs.pop(file_path, None)
        self.modified_texts.pop(file_path, None)
        self.file_timestamps.pop(file_path, None)
        self.search_results.pop(file_path, None)
        self.last_batch_originals.pop(file_path, None)

        if self.current_file == file_path:
            self.current_file = None
            self.current_file_type = None
            self.current_doc = None
            self.current_text_content = ""
            self._clear_text_editor()
            self._clear_formatted_preview()
            self.undo_stack = []
            self.redo_stack = []
            self.update_word_count()
            self.on_tab_change()
            self._update_undo_redo_buttons() # 檔案不存在時也更新按鈕狀態

        self.update_file_tree_display()
        if hasattr(self, 'entry_search'):
            self.search_results = {fp: count for fp, count in self.search_results.items() if fp in self.files}
            self. _update_search_results_ui(self.search_results, self.entry_search.get())

    def _perform_file_load_in_background(self, file_path):
        try:
            ext = os.path.splitext(file_path)[1].lower().lstrip('.')
            loaded_plain_text_content = ""
            loaded_doc_object = None

            if ext == 'txt':
                if file_path in self.modified_texts:
                    loaded_plain_text_content = self.modified_texts[file_path]
                else:
                    # 嘗試多種常見編碼來讀取 TXT 檔案 
                    tried_encodings = ['utf-8', 'gbk', 'big5', 'cp950', 'cp1252', 'latin-1', 'gb2312']
                    loaded_success = False
                    for encoding in tried_encodings:
                        try:
                            with open(file_path, 'r', encoding=encoding) as f:
                                loaded_plain_text_content = f.read()
                            loaded_success = True
                            # 如果成功載入，通知訊息框使用了哪種編碼
                            self.root.after(0, lambda f=file_path, enc=encoding: self._add_message("file_loaded_with_encoding_msg", filename=os.path.basename(f), encoding=enc))
                            break # 成功後就跳出迴圈
                        except UnicodeDecodeError:
                            # 如果當前編碼失敗，繼續嘗試下一個
                            continue
                        except Exception as e_open_txt:
                            # 其他讀取錯誤 (例如檔案不存在或權限問題)
                            self.root.after(0, lambda f=file_path, e=e_open_txt: self._add_message("load_txt_general_error_bg", filename=os.path.basename(f), error=str(e_open_txt)))
                            loaded_plain_text_content = self.tr("file_load_fail_ui_text") # 顯示通用錯誤訊息
                            break # 其他錯誤直接跳出
                    
                    if not loaded_success:
                        # 如果所有嘗試的編碼都失敗了
                        self.root.after(0, lambda f=file_path: self._add_message("load_txt_encoding_fail_bg", filename=os.path.basename(f)))
                        loaded_plain_text_content = self.tr("file_load_fail_ui_text") # 顯示通用錯誤訊息

            elif ext == 'docx':
                if file_path in self.modified_docs and isinstance(self.modified_docs[file_path], docx.document.Document):
                    loaded_doc_object = self.modified_docs[file_path]
                else:
                    try:
                        loaded_doc_object = Document(file_path)
                    except Exception as e_load_docx:
                        error_str = str(e_load_docx)
                        self.root.after(0, lambda filename=os.path.basename(file_path), error=error_str: self._add_message("load_docx_error_bg", filename=filename, error=error))
                        self.root.after(0, self._update_ui_after_load, file_path, self.tr("load_docx_error_ui", error=error_str), None)
                        return
                if loaded_doc_object:
                    loaded_plain_text_content = "\n".join([paragraph.text for paragraph in loaded_doc_object.paragraphs]).strip()
                    if file_path in self.modified_texts:
                        loaded_plain_text_content = self.modified_texts[file_path]
                else:
                    loaded_plain_text_content = self.tr("extract_docx_text_error")
            else:
                self.root.after(0, lambda filename=os.path.basename(file_path): self._add_message("unsupported_file_type_msg", filename=filename))
                self.root.after(0, self._update_ui_after_load, file_path, self.tr("unsupported_file_type_msg", filename=os.path.basename(file_path)), None)
                return
            try:
                self.file_timestamps[file_path] = os.path.getmtime(file_path)
            except Exception as e_ts:
                print(f"載入後無法更新檔案時間戳記 {file_path}: {e_ts}")
                if file_path not in self.file_timestamps: self.file_timestamps[file_path] = 0
            self.root.after(0, self._update_ui_after_load, file_path, loaded_plain_text_content, loaded_doc_object if ext == 'docx' else None)
        except FileNotFoundError:
             self.root.after(0, self._handle_file_not_exist_on_load, file_path)
        except Exception as e:
             error_str_generic = str(e)
             self.root.after(0, lambda filename=os.path.basename(file_path), error=error_str_generic: self._add_message("file_load_error_bg", filename=filename, error=error))
             self.root.after(0, self._update_ui_after_load, file_path, self.tr("file_load_fail_ui_text"), None)

    def _update_ui_after_load(self, file_path, plain_content, doc_object):
        if not (hasattr(self, 'text_content') and self.text_content.winfo_exists() and \
                hasattr(self, 'formatted_preview_text') and self.formatted_preview_text.winfo_exists()):
            return
        self._is_undoing_redoing = True
        self._clear_text_editor()
        self._clear_formatted_preview()

        is_load_error = False
        error_key_templates_for_check = ["load_docx_error_ui", "unsupported_file_type_msg", "file_load_fail_ui_text", "extract_docx_text_error"]
        for err_key_template in error_key_templates_for_check:
            base_error_msg_translated = ""
            if err_key_template == "load_docx_error_ui":
                 base_error_msg_translated = self.tr(err_key_template, error="").split("\n")[0]
            elif err_key_template == "unsupported_file_type_msg":
                 base_error_msg_translated = self.tr(err_key_template, filename="")
            else:
                 base_error_msg_translated = self.tr(err_key_template)

            if plain_content == base_error_msg_translated or \
               (err_key_template == "load_docx_error_ui" and plain_content.startswith(base_error_msg_translated.split(":")[0] + ":")):
                 is_load_error = True
                 break

        if file_path and self.current_file == file_path and not is_load_error:
            self.text_content.configure(state="normal")
            self.text_content.delete("1.0", tk.END)
            self.text_content.insert("1.0", plain_content)
            self.formatted_preview_text.configure(state="normal")
            self.formatted_preview_text.delete("1.0", tk.END)
            self.formatted_preview_text.insert("1.0", plain_content)
            self.formatted_preview_text.configure(state="disabled")
            self.undo_stack = [plain_content]
            self.redo_stack = []
            self.current_text_content = plain_content
            if self.current_file_type == 'docx':
                self.current_doc = doc_object
            else:
                self.current_doc = None
            for item_id in self.file_tree.get_children():
                if self.file_tree.exists(item_id) and self.file_tree.set(item_id, "filename") == file_path:
                    if self.file_tree.exists(item_id):
                        self.file_tree.selection_set(item_id)
                        self.file_tree.focus(item_id)
                        self.file_tree.see(item_id)
                        # 確保選中的項目仍然應用 selected_item_tag
                        current_tags = list(self.file_tree.item(item_id, 'tags'))
                        if "selected_item_tag" not in current_tags:
                            current_tags.append("selected_item_tag")
                            self.file_tree.item(item_id, tags=tuple(current_tags))
                    break
        elif file_path :
            self.text_content.configure(state="normal")
            self.text_content.delete("1.0", tk.END)
            self.text_content.insert("1.0", plain_content)
            self.text_content.configure(state="disabled")
            self.formatted_preview_text.configure(state="normal")
            self.formatted_preview_text.delete("1.0", tk.END)
            self.formatted_preview_text.insert("1.0", plain_content)
            self.formatted_preview_text.configure(state="disabled")
        else:
            self.undo_stack = []
            self.redo_stack = []
            self.current_text_content = ""
            self.current_doc = None

        self._is_undoing_redoing = False
        self.update_word_count()
        self.on_tab_change()
        self.update_file_tree_display()
        self._update_undo_redo_buttons() # 載入文件後，更新撤銷/重做按鈕狀態


    def reload_current_file_from_disk(self):
        if not self.current_file:
            messagebox.showwarning(self.tr("info_title"), self.tr("no_current_file_warning"))
            return
        if not os.path.exists(self.current_file):
            self._handle_file_not_exist_on_load(self.current_file)
            return

        if self.current_file in self.modified_files:
            confirm_reload = messagebox.askyesnocancel(
                self.tr("reload_confirm_title"),
                self.tr("reload_confirm_message", filename=os.path.basename(self.current_file))
            )
            if not confirm_reload:
                return

        self.modified_files.discard(self.current_file)
        self.modified_docs.pop(self.current_file, None)
        self.modified_texts.pop(self.current_file, None)

        self._add_message("reloading_from_disk_msg", filename=os.path.basename(self.current_file))
        self.load_file_content(self.current_file)

    def search_text(self):
        if not (hasattr(self, 'entry_search') and hasattr(self, 'tree')):
            self._add_message(self.tr("components_not_initialized_error"))
            return
        search_term = self.entry_search.get()
        if not search_term:
            self.search_results = {}
            if hasattr(self, 'tree'): self.tree.delete(*self.tree.get_children())
            self._add_message(self.tr("please_enter_search_term_msg"))
            self._highlight_active_editor("")
            self.update_file_tree_display() # 清空搜尋結果後，也需要更新文件列表的顏色
            return

        selected_file_paths = [fp for fp, selected in self.selected_files.items() if selected and fp in self.files]
        if not selected_file_paths:
            messagebox.showwarning(self.tr("warning_title"), self.tr("please_select_files_to_search_msg"))
            self.search_results = {}
            if hasattr(self, 'tree'): self.tree.delete(*self.tree.get_children())
            self._add_message(self.tr("please_select_files_to_search_msg"))
            self._highlight_active_editor("")
            self.update_file_tree_display() # 沒有選擇文件，也要更新文件列表的顏色
            return

        self.search_results = {}
        for i in self.tree.get_children(): self.tree.delete(i)
        self._add_message("search_started_msg", count=len(selected_file_paths), search_term=search_term)
        self._highlight_active_editor("") # 每次新搜尋前清空高亮

        threading.Thread(target=self._perform_search_in_background, args=(selected_file_paths, search_term), daemon=True).start()

    def _perform_search_in_background(self, file_paths, search_term):
        results = {}
        for file_path in file_paths:
            if not os.path.exists(file_path):
                self.root.after(0, lambda filename=os.path.basename(file_path): self._add_message("search_error_file_not_found_msg", filename=filename))
                continue
            try:
                content_to_search = ""
                ext = os.path.splitext(file_path)[1].lower().lstrip('.')
                if ext == 'txt':
                    if file_path in self.modified_texts:
                        content_to_search = self.modified_texts[file_path]
                    else:
                        # 搜尋時也應該使用多編碼載入邏輯，保持一致性
                        tried_encodings = ['utf-8', 'gbk', 'big5', 'cp950', 'cp1252', 'latin-1', 'gb2312']
                        loaded_success_search = False
                        for encoding in tried_encodings:
                            try:
                                with open(file_path, 'r', encoding=encoding) as f_search:
                                    content_to_search = f_search.read()
                                loaded_success_search = True
                                break
                            except UnicodeDecodeError:
                                continue
                            except Exception as e_open_txt_search:
                                # 這裡不發送訊息到 UI，因為是背景搜尋
                                print(f"警告: 背景搜尋 TXT 檔案 '{os.path.basename(file_path)}' 時發生讀取錯誤: {e_open_txt_search}")
                                break
                        if not loaded_success_search:
                            print(f"警告: 背景搜尋 TXT 檔案 '{os.path.basename(file_path)}' 失敗: 所有編碼都無法讀取。")
                            continue # 跳過此檔案的搜尋

                elif ext == 'docx':
                    doc_for_search = None
                    if file_path in self.modified_docs and isinstance(self.modified_docs.get(file_path), docx.document.Document):
                        doc_for_search = self.modified_docs[file_path]
                    else:
                        doc_for_search = Document(file_path)

                    if doc_for_search:
                       content_to_search = "\n".join([paragraph.text for paragraph in doc_for_search.paragraphs])
                    else:
                        content_to_search = ""
                else: continue

                matches = len(re.findall(re.escape(search_term), content_to_search, re.IGNORECASE))
                if matches > 0:
                    results[file_path] = matches
            except Exception as e:
                self.root.after(0, lambda filename=os.path.basename(file_path), error=str(e): self._add_message("search_file_failed_msg", filename=filename, error=error))
        self.root.after(0, self._update_search_results_ui, results, search_term)

    def _update_search_results_ui(self, results, search_term):
        if not (hasattr(self, 'tree') and self.tree.winfo_exists()): return
        self.search_results = results
        for i in self.tree.get_children(): self.tree.delete(i)
        if not results:
            self._add_message("search_completed_no_match_msg", search_term=search_term)
        else:
            sorted_results = sorted(results.items(), key=lambda item: item[1], reverse=True)
            for file_path, count in sorted_results:
                base_filename = os.path.basename(file_path)
                tags_sr = ("red",) # 預設標籤，如果檔案是修改過的
                if file_path not in self.modified_files: # 如果文件沒有被修改，就用藍色表示搜尋結果
                    tags_sr = ("blue",)
                else: # 如果文件同時被修改又有搜尋結果，用綠色表示
                    tags_sr = ("green",)

                self.tree.insert("", "end", values=(base_filename, count), iid=file_path, tags=tags_sr)
                self.tree.set(file_path, "filename", file_path)

            self._add_message("search_completed_matches_found_msg", search_term=search_term, count=len(results))
        self.update_file_tree_display()
        if self.current_file:
            self._highlight_active_editor(search_term)

    def on_tree_click(self, event):
        if not (hasattr(self, 'tree') and self.tree.winfo_exists()): return
        item_id = self.tree.identify_row(event.y)
        if not item_id or not self.tree.exists(item_id): return

        full_file_path = self.tree.set(item_id, "filename")
        if not full_file_path or full_file_path not in self.files:
             self._add_message("file_not_in_list_warning", filename=os.path.basename(full_file_path if full_file_path else item_id))
             return

        proceed_with_switch_sr = True
        if self.current_file and self.current_file in self.modified_files:
            if self.auto_save_var.get():
                 self._add_message("auto_saving_modified_file_msg", filename=os.path.basename(self.current_file))
                 save_success_sr = self.save_file(file_path_to_save=self.current_file)
                 if not save_success_sr: proceed_with_switch_sr = False
            else:
                confirm_save_sr = messagebox.askyesnocancel(
                    self.tr("save_changes_on_switch_title"),
                    self.tr("save_changes_on_switch_message", filename=os.path.basename(self.current_file))
                )
                if confirm_save_sr is True:
                    save_success_sr = self.save_file()
                    if not save_success_sr: proceed_with_switch_sr = False
                elif confirm_save_sr is False:
                    self.modified_files.discard(self.current_file)
                    self.modified_docs.pop(self.current_file, None)
                    self.modified_texts.pop(self.current_file, None)
                    self.update_file_tree_display()
                else:
                    proceed_with_switch_sr = False

        if proceed_with_switch_sr:
            self.load_file_content(full_file_path)
            self.root.after(200, lambda: self._highlight_active_editor(self.entry_search.get()))

            if hasattr(self, 'tab_view') and hasattr(self.tab_view, "_segmented_button") and self.tab_view._segmented_button is not None:
                try:
                    first_tab_name_translated = self.tr("tab_text_edit")
                    if first_tab_name_translated in self.tab_view._tab_dict:
                        self.tab_view.set(first_tab_name_translated)
                    else:
                        if self.tab_view._segmented_button.cget("values"):
                             self.tab_view.set(self.tab_view._segmented_button.cget("values")[0])
                    self.on_tab_change()
                except Exception as e_tab_switch:
                    print(f"從搜尋結果切換tab時發生錯誤: {e_tab_switch}")

    def execute_replace_action(self, mode):
        replace_rules = []
        if hasattr(self, 'replace_tree'):
            for item_id in self.replace_tree.get_children():
                if self.replace_tree.exists(item_id):
                    values = self.replace_tree.item(item_id)['values']
                    if values and len(values) == 2:
                        replace_rules.append((values[0], values[1]))

        if not replace_rules:
            messagebox.showwarning(self.tr("warning_title"), self.tr("add_replace_rules_warning"))
            return

        selected_file_paths = [fp for fp, selected in self.selected_files.items() if selected and fp in self.files]
        if not selected_file_paths:
            messagebox.showwarning(self.tr("warning_title"), self.tr("at_least_one_file_error"))
            return

        auto_save_info_line = self.tr("auto_save_enabled_info_line") if self.auto_save_var.get() else ""

        if mode == "plain_text_batch":
            confirm_msg = self.tr("confirm_plain_batch_message",
                                  count=len(selected_file_paths),
                                  auto_save_info=auto_save_info_line)
            if messagebox.askyesno(self.tr("confirm_plain_batch_title"), confirm_msg):
                self._add_message("starting_plain_batch_msg", count=len(selected_file_paths))
                threading.Thread(target=self._perform_replace_in_background, args=(selected_file_paths, replace_rules, "plain"), daemon=True).start()
            else:
                self._add_message("user_cancelled_plain_batch_msg")
        elif mode == "formatted_batch":
            confirm_msg = self.tr("confirm_formatted_batch_message",
                                  count=len(selected_file_paths),
                                  auto_save_info=auto_save_info_line)
            if messagebox.askyesno(self.tr("confirm_formatted_batch_title"), confirm_msg):
                self._add_message("starting_formatted_batch_msg", count=len(selected_file_paths))
                threading.Thread(target=self._perform_replace_in_background, args=(selected_file_paths, replace_rules, "formatted"), daemon=True).start()
            else:
                self._add_message("user_cancelled_formatted_batch_msg")
        else:
            self._add_message("unknown_replace_mode_error", mode=mode)
            print(f"錯誤：未知的取代模式 {mode}")


    def _perform_replace_in_background(self, file_paths, rules, mode):
        self.last_batch_originals.clear()
        # 備份原始檔案內容
        for fp_orig in file_paths:
            if not os.path.exists(fp_orig): continue
            ext_orig = os.path.splitext(fp_orig)[1].lower().lstrip('.')
            try:
                if ext_orig == 'txt':
                    with open(fp_orig, 'r', encoding='utf-8') as f_orig:
                        self.last_batch_originals[fp_orig] = {"type": "txt", "content": f_orig.read()}
                elif ext_orig == 'docx':
                    # 備份 DOCX 時，複製其 Document 物件以避免修改原始檔案的記憶體狀態
                    original_doc = Document(fp_orig)
                    self.last_batch_originals[fp_orig] = {"type": "docx", "content": original_doc}
            except Exception as e_backup:
                 self.root.after(0, lambda f=fp_orig, e=e_backup: self._add_message(f"無法備份原始檔案 {os.path.basename(f)}: {e}"))

        modified_count_actual = 0
        for index, file_path in enumerate(file_paths):
            # 每處理 10 個檔案或到達總數時更新進度訊息
            if (index + 1) % 10 == 0 or (index + 1) == len(file_paths):
                self.root.after(0, lambda f=file_path, i=index, t=len(file_paths), m=mode: self._update_batch_progress_message(i + 1, t, os.path.basename(f), m))

            if not os.path.exists(file_path):
                self.root.after(0, lambda f=file_path: self._add_message("batch_error_file_not_found_msg", filename=os.path.basename(f)))
                continue
            try:
                file_ext = os.path.splitext(file_path)[1].lower().lstrip('.')

                if file_ext == 'txt':
                    current_content = ""
                    if file_path in self.modified_texts:
                        current_content = self.modified_texts[file_path]
                    else:
                        # 批次取代時也使用多編碼載入邏輯，保持一致性
                        tried_encodings_batch = ['utf-8', 'gbk', 'big5', 'cp950', 'cp1252', 'latin-1', 'gb2312']
                        loaded_success_batch = False
                        for encoding_batch in tried_encodings_batch:
                            try:
                                with open(file_path, 'r', encoding=encoding_batch) as f_batch:
                                    current_content = f_batch.read()
                                loaded_success_batch = True
                                break
                            except UnicodeDecodeError:
                                continue
                            except Exception as e_open_txt_batch:
                                # 這裡不發送訊息到 UI，因為是背景批次處理
                                print(f"警告: 背景批次處理 TXT 檔案 '{os.path.basename(file_path)}' 時發生讀取錯誤: {e_open_txt_batch}")
                                break
                        if not loaded_success_batch:
                            print(f"警告: 背景批次處理 TXT 檔案 '{os.path.basename(file_path)}' 失敗: 所有編碼都無法讀取。")
                            self.root.after(0, lambda f=file_path: self._add_message("batch_replace_file_error_msg", filename=os.path.basename(f), error=self.tr("load_txt_encoding_fail_bg", filename=os.path.basename(f))))
                            continue # 跳過此檔案的批次處理

                    original_content_before_this_file_replace = current_content
                    text_changed_in_file = False
                    for search_text, replace_text in rules:
                        new_content = current_content.replace(search_text, replace_text)
                        if new_content != current_content:
                            text_changed_in_file = True
                        current_content = new_content

                    if text_changed_in_file:
                        self.modified_texts[file_path] = current_content
                        self.modified_files.add(file_path)
                        self.root.after(0, lambda f=file_path: self._add_message("txt_modified_in_memory_msg", filename=os.path.basename(f)))
                        modified_count_actual += 1
                        if self.auto_save_var.get():
                            self._save_file_content_to_disk(file_path, "txt", self.modified_texts[file_path])
                            # 自動儲存成功後，從記憶體中移除修改標記和內容
                            self.modified_files.discard(file_path)
                            self.modified_texts.pop(file_path, None)

                elif file_ext == 'docx':
                    doc_to_process = None
                    if file_path in self.modified_docs and isinstance(self.modified_docs.get(file_path), docx.document.Document):
                        doc_to_process = self.modified_docs[file_path]
                    else:
                        try:
                            doc_to_process = Document(file_path)
                        except Exception as e_load_docx_batch:
                            self.root.after(0, lambda f=file_path, e=e_load_docx_batch: self._add_message("batch_error_load_docx_fail_msg", filename=os.path.basename(f), error=str(e)))
                            continue

                    doc_changed = False
                    if mode == "plain":
                        full_text = "\n".join([p.text for p in doc_to_process.paragraphs])
                        original_full_text = full_text
                        for search_text, replace_text in rules:
                            full_text = full_text.replace(search_text, replace_text)
                        if full_text != original_full_text:
                            doc_changed = True
                            new_doc_plain = Document()
                            for line in full_text.splitlines():
                                new_doc_plain.add_paragraph(line)
                            self.modified_docs[file_path] = new_doc_plain
                            self.modified_texts[file_path] = full_text

                    elif mode == "formatted":
                        for table in doc_to_process.tables:
                            for row in table.rows:
                                for cell in row.cells:
                                    for paragraph in cell.paragraphs:
                                        original_text = paragraph.text
                                        current_para_text = original_text
                                        for search_text, replace_text in rules:
                                            if search_text in current_para_text:
                                                current_para_text = current_para_text.replace(search_text, replace_text)
                                        if current_para_text != original_text:
                                            paragraph.text = current_para_text
                                            doc_changed = True
                        for paragraph in doc_to_process.paragraphs:
                            # 檢查段落是否在文檔主體中，避免處理表格內的段落兩次
                            if isinstance(paragraph._element, docx.oxml.text.paragraph.CT_P) and paragraph._element.getparent().tag.endswith('body'):
                                original_text = paragraph.text
                                current_para_text = original_text
                                for search_text, replace_text in rules:
                                    if search_text in current_para_text:
                                        current_para_text = current_para_text.replace(search_text, replace_text)
                                if current_para_text != original_text:
                                    paragraph.text = current_para_text
                                    doc_changed = True
                        self.modified_docs[file_path] = doc_to_process
                        updated_plain_text = "\n".join([p.text for p in doc_to_process.paragraphs]).strip()
                        self.modified_texts[file_path] = updated_plain_text

                    if doc_changed:
                        self.modified_files.add(file_path)
                        self.root.after(0, lambda f=file_path: self._add_message("docx_modified_in_memory_msg", filename=os.path.basename(f)))
                        modified_count_actual += 1
                        if self.auto_save_var.get():
                            self._save_file_content_to_disk(file_path, "docx", self.modified_docs[file_path], self.modified_texts[file_path], allow_format_loss=True)
                            # 自動儲存成功後，從記憶體中移除修改標記和內容
                            self.modified_files.discard(file_path)
                            self.modified_docs.pop(file_path, None)
                            self.modified_texts.pop(file_path, None)
                else:
                    self.root.after(0, lambda f=file_path: self._add_message("batch_skip_unsupported_type_msg", filename=os.path.basename(f)))

            except Exception as e_file_proc:
                self.root.after(0, lambda f=file_path, e=e_file_proc: self._add_message("batch_replace_file_error_msg", filename=os.path.basename(f), error=str(e)))
                traceback.print_exc()

        self.root.after(0, self._update_replace_results_ui, modified_count_actual)
        if hasattr(self, 'btn_undo_batch'): self.root.after(0, lambda: self.btn_undo_batch.configure(state="normal"))


    def _update_batch_progress_message(self, current, total, filename, mode_display):
        # 這裡需要一個 "replace_mode_display_plain" 或 "replace_mode_display_formatted" 的翻譯鍵
        # 因為 app_languages.py 中沒有定義，所以先假設會直接顯示 mode_display
        progress_msg = f"批次處理 ({mode_display}) 中： {current}/{total} - {filename}"
        # Only add message for first, every 10th, and last file
        if current == 1 or current % 10 == 0 or current == total:
             self._add_message(progress_msg)


    def _update_replace_results_ui(self, count):
        if count > 0:
            self._add_message("batch_replace_done_modified_msg", count=count)
            # 如果目前檔案被修改了，且沒有在 modified_files 中，則重新載入 (可能被自動儲存)
            if self.current_file and (self.current_file in self.modified_files or
                                      (self.auto_save_var.get() and self.current_file not in self.modified_files)):
                 self.load_file_content(self.current_file)
        else:
            self._add_message("batch_replace_done_no_changes_msg")
        self.update_file_tree_display()
        if hasattr(self, 'entry_search') and self.entry_search.get():
            self.search_text()

    def undo_last_batch_replace(self):
        if not self.last_batch_originals:
            messagebox.showinfo(self.tr("info_title"), self.tr("no_batch_undo_available_info"))
            return

        confirm_msg = self.tr("confirm_undo_batch_message", count=len(self.last_batch_originals))
        if messagebox.askyesno(self.tr("confirm_undo_batch_title"), confirm_msg):
            self._add_message("starting_undo_batch_msg")
            threading.Thread(target=self._perform_undo_batch_in_background, daemon=True).start()
        else:
            self._add_message("user_cancelled_undo_batch_msg")

    def _perform_undo_batch_in_background(self):
        restored_count = 0
        for file_path, original_data in self.last_batch_originals.items():
            try:
                if not os.path.exists(os.path.dirname(file_path)):
                     self.root.after(0, lambda f=file_path: self._add_message(f"無法復原 {os.path.basename(f)}: 目錄不存在。"))
                     continue

                original_type = original_data.get("type")
                original_content = original_data.get("content")

                if original_type == "txt":
                    self.modified_texts[file_path] = original_content
                    self.modified_files.add(file_path)
                    self.modified_docs.pop(file_path, None) # 確保移除 DOCX 相關的記憶體內容
                    self.root.after(0, lambda f=file_path: self._add_message("txt_restored_in_memory_msg", filename=os.path.basename(f)))
                    restored_count += 1
                elif original_type == "docx" and isinstance(original_content, docx.document.Document):
                    self.modified_docs[file_path] = original_content
                    restored_plain_text = "\n".join([p.text for p in original_content.paragraphs]).strip()
                    self.modified_texts[file_path] = restored_plain_text
                    self.modified_files.add(file_path)
                    self.root.after(0, lambda f=file_path: self._add_message("docx_restored_in_memory_msg", filename=os.path.basename(f)))
                    restored_count += 1
                else:
                    self.root.after(0, lambda f=file_path: self._add_message("cannot_undo_unknown_data_msg", filename=os.path.basename(f)))
            except Exception as e_undo:
                self.root.after(0, lambda f=file_path, e=e_undo: self._add_message("undo_file_failed_msg", filename=os.path.basename(f), error=str(e)))

        if restored_count > 0:
            self.root.after(0, lambda c=restored_count: self._add_message("undo_batch_done_restored_msg", count=c))
        else:
            self.root.after(0, lambda: self._add_message("undo_batch_done_no_restore_msg"))

        self.last_batch_originals.clear()
        if hasattr(self, 'btn_undo_batch'): self.root.after(0, lambda: self.btn_undo_batch.configure(state="disabled"))

        # 如果當前開啟的檔案在被復原的檔案列表中，重新載入其內容
        if self.current_file and self.current_file in self.files and self.current_file in self.modified_files:
            self.root.after(0, self.load_file_content, self.current_file)
        # 如果當前檔案在被復原的檔案列表中，但現在已經沒有修改狀態了，也要重新載入以反映磁碟內容 (如果自動儲存是關閉的)
        elif self.current_file and self.current_file in self.files and self.current_file not in self.modified_files:
             self.root.after(0, self.load_file_content, self.current_file)

        self.root.after(0, self.update_file_tree_display)
        self.root.after(0, lambda: self._highlight_active_editor(self.entry_search.get()))
        self.root.after(0, self._update_undo_redo_buttons) # 復原批次操作後，更新撤銷/重做按鈕狀態


    def _save_file_content_to_disk(self, file_path, file_ext, content_or_doc, plain_text_content=None, allow_format_loss=False):
        """將記憶體中的內容或Document物件儲存到磁碟。

        Args:
            file_path (str): 要儲存的檔案路徑。
            file_ext (str): 檔案副檔名 (txt 或 docx)。
            content_or_doc: 如果是 txt 檔案，為純文字內容；如果是 docx 檔案，為 Document 物件。
            plain_text_content (str, optional): 如果是 docx 檔案且從純文字儲存，此為其純文字內容。
            allow_format_loss (bool, optional): 是否允許 docx 儲存時格式遺失。
        Returns:
            bool: 儲存是否成功。
        """
        try:
            if file_ext == 'txt':
                if content_or_doc is None:
                    self._add_message("txt_warn_modified_no_content", filename=os.path.basename(file_path))
                    return False
                with open(file_path, 'w', encoding='utf-8') as f:
                    f.write(content_or_doc)

            elif file_ext == 'docx':
                if isinstance(content_or_doc, docx.document.Document):
                    self._add_message("docx_save_from_memory_formatted_msg", filename=os.path.basename(file_path))
                    content_or_doc.save(file_path)
                elif plain_text_content is not None:
                    if not allow_format_loss:
                        confirm_format_loss = messagebox.askyesno(
                            self.tr("docx_format_loss_warning_title"),
                            self.tr("docx_format_loss_warning_message", filename=os.path.basename(file_path))
                        )
                        if not confirm_format_loss:
                            self._add_message("docx_cancelled_save_format_loss_msg", filename=os.path.basename(file_path))
                            return False

                    self._add_message("docx_save_from_plain_text_msg", filename=os.path.basename(file_path))
                    new_doc = Document()
                    for line in plain_text_content.splitlines():
                        new_doc.add_paragraph(line)
                    new_doc.save(file_path)
                else:
                    self._add_message("docx_warn_modified_no_plain_text", filename=os.path.basename(file_path))
                    return False
            else:
                self._add_message("unsupported_file_type_cannot_save_msg", filename=os.path.basename(file_path))
                return False

            try:
                self.file_timestamps[file_path] = os.path.getmtime(file_path)
            except Exception as e_ts_save:
                 self._add_message("file_save_timestamp_error", filepath=os.path.basename(file_path), error=str(e_ts_save))

            self._add_message("file_save_successful_msg", filename=os.path.basename(file_path))
            return True

        except Exception as e_s:
            self._add_message("file_save_failed_error_message", filename=os.path.basename(file_path), error=str(e_s))
            traceback.print_exc()
            return False

    def save_file(self, file_path_to_save=None, is_batch_save=False, event=None): # Add event parameter
        target_file_path = file_path_to_save if file_path_to_save else self.current_file
        if not target_file_path:
            if not is_batch_save: messagebox.showwarning(self.tr("warning_title"), self.tr("no_file_to_save_warning"))
            return False
        if not os.path.exists(target_file_path):
            if not is_batch_save: messagebox.showerror(self.tr("error_title"), self.tr("file_not_exist_cannot_save_error", filename=os.path.basename(target_file_path)))
            self._add_message("file_not_exist_cannot_save_error", filename=os.path.basename(target_file_path))
            return False
        if target_file_path not in self.modified_files and not is_batch_save:
            if not is_batch_save: self._add_message("file_no_unsaved_changes_msg", filename=os.path.basename(target_file_path))
            return True

        self._add_message("saving_file_msg", filename=os.path.basename(target_file_path))
        file_ext_save = os.path.splitext(target_file_path)[1].lower().lstrip('.')

        content_or_doc = None
        plain_text_content = None

        if file_ext_save == 'txt':
            content_or_doc = self.modified_texts.get(target_file_path)
        elif file_ext_save == 'docx':
            content_or_doc = self.modified_docs.get(target_file_path)
            plain_text_content = self.modified_texts.get(target_file_path)

        success = self._save_file_content_to_disk(target_file_path, file_ext_save, content_or_doc, plain_text_content, allow_format_loss=is_batch_save)

        if success:
            self.modified_files.discard(target_file_path)
            self.modified_docs.pop(target_file_path, None) # 成功儲存後，從記憶體中移除對 docx 物件的引用
            self.modified_texts.pop(target_file_path, None) # 成功儲存後，從記憶體中移除對文本內容的引用

            if self.current_file == target_file_path:
                self.update_word_count()
            self.update_file_tree_display()
        return success

    def save_all_selected_files(self):
        selected_and_modified_files = [
            fp for fp, selected in self.selected_files.items()
            if selected and fp in self.modified_files and fp in self.files
        ]
        if not selected_and_modified_files:
            messagebox.showinfo(self.tr("info_title"), self.tr("no_selected_modified_files_to_save_info"))
            return

        confirm_msg = self.tr("confirm_save_selected_message", count=len(selected_and_modified_files))
        if messagebox.askyesno(self.tr("confirm_save_selected_title"), confirm_msg):
            self._add_message("saving_all_selected_files_msg", count=len(selected_and_modified_files))

            saved_count = 0
            failed_count = 0
            for file_path in selected_and_modified_files:
                try:
                    if self.save_file(file_path_to_save=file_path, is_batch_save=True):
                        saved_count += 1
                    else:
                        failed_count += 1
                except Exception as e_batch_save:
                     failed_count +=1
                     self._add_message("batch_save_unexpected_error_msg", filename=os.path.basename(file_path), error=str(e_batch_save))
                     traceback.print_exc()

            if failed_count > 0:
                 messagebox.showwarning(self.tr("save_partial_fail_warning_title"),
                                        self.tr("save_partial_fail_warning_message", saved_count=saved_count, failed_count=failed_count))
                 self._add_message("all_selected_saved_done_msg", saved_count=saved_count, failed_count=failed_count)
            else:
                 messagebox.showinfo(self.tr("save_all_done_info_title"), self.tr("save_all_done_info_message", saved_count=saved_count))
                 self._add_message("all_selected_saved_done_msg", saved_count=saved_count, failed_count=0)

            self.update_file_tree_display()
            if self.current_file in selected_and_modified_files and self.current_file not in self.modified_files:
                self.update_word_count()
            self._highlight_active_editor(self.entry_search.get())
        else:
            self._add_message("user_cancelled_save_all_selected_msg")

    def delete_selected_files(self):
        selected_to_remove = [fp for fp, selected in self.selected_files.items() if selected and fp in self.files]
        if not selected_to_remove:
            messagebox.showwarning(self.tr("warning_title"), self.tr("select_files_to_remove_from_list_warning"))
            return

        unsaved_in_selection = [fp for fp in selected_to_remove if fp in self.modified_files]
        unsaved_warning_line = ""
        if unsaved_in_selection:
            unsaved_basenames_short = [os.path.basename(f) for f in unsaved_in_selection[:5]]
            if len(unsaved_in_selection) > 5: unsaved_basenames_short.append("...")
            file_list_str_remove = "\n".join(unsaved_basenames_short)
            unsaved_warning_line = self.tr("unsaved_files_warning_prefix_line", count=len(unsaved_in_selection), file_list=file_list_str_remove)

        confirm_msg_remove = self.tr("confirm_remove_files_message", count=len(selected_to_remove), unsaved_warning=unsaved_warning_line)
        if messagebox.askyesno(self.tr("confirm_remove_files_title"), confirm_msg_remove):
            removed_count = 0
            for file_path in selected_to_remove:
                try:
                    if file_path in self.files: self.files.remove(file_path)
                    self.selected_files.pop(file_path, None)
                    self.modified_files.discard(file_path)
                    self.modified_docs.pop(file_path, None)
                    self.modified_texts.pop(file_path, None)
                    self.file_timestamps.pop(file_path, None)
                    self.search_results.pop(file_path, None)
                    self.last_batch_originals.pop(file_path, None)

                    # 如果移除的是當前開啟的檔案，清空編輯器並重置狀態
                    if self.current_file == file_path:
                        self.current_file = None
                        self.current_file_type = None
                        self.current_doc = None
                        self.current_text_content = ""
                        self._clear_text_editor()
                        self._clear_formatted_preview()
                        self.undo_stack = []
                        self.redo_stack = []
                        self.update_word_count()
                        self.on_tab_change()
                        self._update_undo_redo_buttons() # 移除當前檔案後，更新按鈕狀態

                    self._add_message("removed_from_list_msg", filename=os.path.basename(file_path))
                    removed_count += 1
                except Exception as e_remove:
                     self._add_message("remove_file_failed_msg", filename=os.path.basename(file_path), error=str(e_remove))

            if removed_count > 0:
                 self._add_message("files_removed_from_list_msg", count=removed_count)
                 messagebox.showinfo(self.tr("files_removed_done_info_title"), self.tr("files_removed_done_info_message", count=removed_count))

            self.update_file_tree_display()
            self.update_select_all_state()
            if hasattr(self, 'entry_search'):
                current_search_term = self.entry_search.get()
                self.search_results = {fp: count for fp, count in self.search_results.items() if fp in self.files}
                self._update_search_results_ui(self.search_results, current_search_term)
        else:
            self._add_message("user_cancelled_remove_files_msg")

    def select_all(self, event=None):
        widget = event.widget
        if widget == self.file_tree:
            self.select_all_var.set(not self.select_all_var.get())
            self.toggle_select_all()
            return "break"
        elif isinstance(widget, (ctk.CTkEntry, ctk.CTkTextbox, ttk.Entry)):
            self.select_all_text(event)
            return "break"
        return None

# --- 主程式執行 ---
if __name__ == "__main__":
    root = TkinterDnD.Tk()
    app = WordEditorApp(root)
    root.mainloop()
